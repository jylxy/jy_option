#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build a Feishu-import friendly DOCX report for S1 backtest analysis.

The script intentionally keeps the source file UTF-8 and reads/writes all text
with explicit encodings. It converts a Markdown report into DOCX, then appends
standard chart chapters with embedded PNGs and deterministic commentary.
"""

from __future__ import annotations

import argparse
import csv
import math
import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence

try:
    import pandas as pd
except ImportError:  # pragma: no cover - script is expected to run in Codex runtime.
    pd = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx is required. Use the bundled Codex Python runtime.") from exc


CHART_SPECS = [
    {
        "file": "01_nav_drawdown.png",
        "title": "净值曲线与最大回撤",
        "read": "这张图同时看收益斜率和回撤深度，重点观察收益是否平滑、回撤是否集中在少数事件日。",
        "meaning": "如果净值主要靠长时间小幅爬升、少数日期大幅回撤，说明策略本质仍是卖尾部保险，需要继续控制跳跃和升波阶段。",
    },
    {
        "file": "02_margin_positions.png",
        "title": "保证金使用率与持仓数量",
        "read": "这张图看仓位是否真正打满、持仓品种和合约数是否稳定，以及回撤期是否被动降仓。",
        "meaning": "若保证金稳定但收益不高，问题通常不在仓位太低，而在单位风险的 theta 质量、gamma/vega 吞噬和合约选择效率。",
    },
    {
        "file": "03_greeks_timeseries.png",
        "title": "组合 Greeks 时序",
        "read": "这张图看 cash delta、cash gamma、cash vega 是否长期偏向某一方向，以及尾部日期前是否已经暴露过高。",
        "meaning": "卖权策略可以接受 short gamma 和 short vega，但必须确认这些暴露获得了足够 theta 补偿，否则就是低质量卖权。",
    },
    {
        "file": "04_pnl_attribution.png",
        "title": "PnL 归因时序",
        "read": "这张图拆解 delta、gamma、theta、vega 和 residual 的累计贡献，判断策略到底赚的是 carry 还是方向。",
        "meaning": "目标画像应是 theta 稳定为正、vega 在降波段贡献为正、gamma 亏损受控；如果 vega 长期为负，说明入场环境或 IV 口径仍需优化。",
    },
    {
        "file": "05_daily_pnl_tail.png",
        "title": "日度 PnL 左尾",
        "read": "这张图看最差日期是否集中在升波、跳空或系统性事件里，以及单日亏损是否吞掉多月 theta。",
        "meaning": "卖方策略的真实质量不由胜率决定，而由最差日、最差周和回撤修复速度决定。",
    },
    {
        "file": "06_premium_pc_structure.png",
        "title": "权利金和 P/C 结构",
        "read": "这张图看 Put/Call 暴露是否稳定、是否出现单侧拥挤，以及 open premium 与 liability 是否同步变化。",
        "meaning": "P/C 极端偏离会把卖波策略变成隐含方向策略，后续应加入趋势置信度、单侧预算和再平衡规则。",
    },
    {
        "file": "07_vol_regime_exposure.png",
        "title": "波动状态暴露",
        "read": "这张图看仓位是否集中在 falling、low、normal、high、post-stop 等状态，以及不同状态下的收益质量。",
        "meaning": "如果 falling 状态仓位不足，策略会错过最适合卖方的降波收益；如果 low-vol 状态仓位过重，则容易在波动切换时受伤。",
    },
    {
        "file": "08_calendar_returns.png",
        "title": "月度收益热力图",
        "read": "这张图看收益是否跨年份稳定，还是集中在少数月份；同时定位关键回撤月份。",
        "meaning": "稳定卖权策略不需要每月赚钱，但不能让少数月份反复吃掉全年 theta。",
    },
    {
        "file": "09_product_share_top10.png",
        "title": "Top10 品种持仓占比",
        "read": "这张图看真实仓位是否被少数品种主导，以及品种集中度是否随时间升高。",
        "meaning": "全品种扫描不等于真正分散，如果成交和 Delta 规则导致实际只集中在少数商品，组合仍然有板块和相关性尾部风险。",
    },
    {
        "file": "10_order_action_summary.png",
        "title": "订单动作汇总",
        "read": "这张图看开仓、到期、止损、期末持仓等动作比例，判断收益主要来自持有到期还是频繁止损。",
        "meaning": "若止损占比过高，说明入场筛选或异常报价过滤不足；若几乎全到期，则要重点检查到期前 gamma 风险。",
    },
    {
        "file": "11_close_event_timeline.png",
        "title": "平仓事件时间线",
        "read": "这张图看止损是否簇集、是否集中在回撤月份，以及止损后策略是否过快重开同类风险。",
        "meaning": "止损簇集是卖方策略的核心风险信号，后续冷却期和波动回落确认应围绕它设计。",
    },
    {
        "file": "12_daily_open_premium_vega_quality.png",
        "title": "每日新收权利金与 Vega 吞噬率",
        "read": "这张图把每日新开仓权利金、权利金占 NAV、累计权利金、theta、vega/gamma 损耗放在一起看，核心问题不是卖了多少，而是收来的权利金是否留得住。",
        "meaning": "S1 后续主线应从放大开仓切到控制 vega 和提高权利金留存率；如果毛权利金足够高但 vega loss / gross premium 过大，说明策略质量仍不合格。",
    },
    {
        "file": "13_tail_product_side_contribution.png",
        "title": "尾部日产品与方向贡献",
        "read": "这张图把最差 NAV 日期里的平仓实现损益按产品和 Call/Put 拆开，用来定位尾部亏损是否集中在少数品种或单侧方向。",
        "meaning": "如果尾部亏损集中在少数产品/方向，后续应优先加入板块、相关性、单侧预算和异常报价复核；注意该图是平仓实现损益，不等同于全部盯市浮亏。",
    },
    {
        "file": "14_vega_quality_by_bucket.png",
        "title": "分桶 Vega 质量",
        "read": "这张图按波动状态和方向比较权利金留存率、premium/vega proxy 与开仓权利金规模，判断哪些环境下卖出的保险费更容易留下来。",
        "meaning": "好的 bucket 应该同时具备较高留存率和合理 premium/vega proxy；如果 high/rising bucket 权利金厚但留存差，不能因为权利金高就加仓。",
    },
    {
        "file": "15_stop_slippage_distribution.png",
        "title": "止损滑点分布",
        "read": "这张图观察止损平仓时的成交滑点分布、滑点占原始权利金比例，以及滑点集中在哪些产品。",
        "meaning": "卖权策略的尾部损失会被止损滑点放大；若滑点在低流动性或跳价产品中集中，实盘应加入延迟确认、异常报价过滤和产品级流动性降权。",
    },
    {
        "file": "16_pc_funnel.png",
        "title": "Put/Call 权利金漏斗",
        "read": "这张图比较 Put 和 Call 从开仓手数、开仓权利金、已平仓留存权利金到止损占比的漏斗，判断 P/C 偏移是否真正贡献了净收益。",
        "meaning": "P/C 不是越偏越好；如果某一侧收得多但留存率低或止损占比高，说明该侧卖权质量差，后续应降低该侧预算或要求更强趋势/降波确认。",
    },
    {
        "file": "compare_01_nav_relative_to_b0.png",
        "title": "相对 B0 的净值与超额净值",
        "read": "这张图把候选版本和 B0 放在同一共同区间里比较，重点看超额净值是否持续，而不是只看最终点位是否更高。",
        "meaning": "如果超额净值来自少数跳升日，说明改动可能只是路径运气；如果超额净值稳定抬升，才更像结构性改善。",
    },
    {
        "file": "compare_02_drawdown_relative_to_b0.png",
        "title": "相对 B0 的回撤路径",
        "read": "这张图比较候选版本和 B0 的回撤深度、恢复速度和回撤同步性。",
        "meaning": "卖权版本升级不能只提高收益，必须确认回撤没有被隐性放大，尤其要看升波阶段是否比 B0 更脆弱。",
    },
    {
        "file": "compare_03_margin_position_relative_to_b0.png",
        "title": "相对 B0 的保证金与持仓厚度",
        "read": "这张图看候选版本是否只是用了更多保证金、更多合约或更多手数来换收益。",
        "meaning": "如果收益改善主要来自仓位更厚，需要继续用单位保证金收益和尾部亏损验证，而不能直接视为 alpha。",
    },
    {
        "file": "compare_04_greek_attribution_relative_to_b0.png",
        "title": "相对 B0 的 Greek 归因差异",
        "read": "这张图比较 theta、vega、gamma、delta、residual 的累计差异，是判断卖波质量是否改善的核心图。",
        "meaning": "理想升级应增加高质量 theta，并改善或至少不显著恶化 vega/gamma；若 NAV 更好但 vega/gamma 更差，要标记为风险放大。",
    },
    {
        "file": "compare_05_daily_pnl_tail_relative_to_b0.png",
        "title": "相对 B0 的日度左尾风险",
        "read": "这张图比较日度 PnL 分布、左尾分位和最差日期，判断收益改善是否伴随更厚左尾。",
        "meaning": "卖方策略最怕用更多小赢换更大单日亏损；左尾没有同步改善时，收益提升的可靠性要打折。",
    },
    {
        "file": "compare_06_pc_structure_relative_to_b0.png",
        "title": "相对 B0 的 P/C 与方向结构",
        "read": "这张图比较 Put/Call 手数、Call 占比和 P/C 偏离，判断候选版本是否变成方向仓。",
        "meaning": "B0 是基准卖权结构，候选版本如果通过 P/C 偏移获得收益，需要区分是合理趋势适配还是方向押注。",
    },
    {
        "file": "compare_07_product_exposure_relative_to_b0.png",
        "title": "相对 B0 的品种暴露差异",
        "read": "这张图比较 Top 品种占比和产品贡献差异，判断候选版本是否更可交易、更分散，还是更集中。",
        "meaning": "流动性排序或持仓排序版本尤其要看这张图：如果收益来自少数品种集中暴露，后续必须叠加板块和相关性约束。",
    },
    {
        "file": "compare_08_stop_cluster_relative_to_b0.png",
        "title": "相对 B0 的止损簇集差异",
        "read": "这张图比较止损次数、止损金额和簇集路径，判断候选版本是否减少了真实尾部事件。",
        "meaning": "止损减少如果来自更好的流动性或异常报价过滤，是高质量改善；如果只是风险尚未暴露，则需要更长样本确认。",
    },
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build S1 backtest DOCX report with embedded charts.")
    parser.add_argument("--tag", required=True, help="Backtest tag, e.g. s1_b0_standard_stop25_allprod_2022_latest.")
    parser.add_argument("--markdown", required=True, type=Path, help="Source Markdown report.")
    parser.add_argument("--analysis-dir", type=Path, help="Analysis directory. Default: output/analysis_<tag>.")
    parser.add_argument("--output-docx", type=Path, help="Output DOCX path.")
    parser.add_argument("--title", default=None, help="Document title override.")
    parser.add_argument("--repo-root", type=Path, default=Path.cwd(), help="Repository root.")
    parser.add_argument("--skip-existing-chart-section", action="store_true", help="Do not append chart chapters.")
    parser.add_argument(
        "--skip-markdown-images",
        action="store_true",
        help="Ignore Markdown image syntax. Use this when Feishu import should rely on the explicit chart section.",
    )
    return parser.parse_args()


def fmt_num(value: object, decimals: int = 2) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "N/A"
    if math.isnan(number):
        return "N/A"
    return f"{number:,.{decimals}f}"


def fmt_pct(value: object, decimals: int = 2) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "N/A"
    if math.isnan(number):
        return "N/A"
    return f"{number * 100:.{decimals}f}%"


def safe_read_csv(path: Path) -> Optional["pd.DataFrame"]:
    if pd is None or not path.exists():
        return None
    try:
        return pd.read_csv(path)
    except Exception:
        return None


def read_metrics_csv(path: Path) -> Dict[str, str]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    return rows[0] if rows else {}


def read_comparison_summary(path: Path) -> Dict[str, str]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    result: Dict[str, str] = {}
    for row in rows:
        metric = str(row.get("metric", "")).strip()
        if not metric:
            continue
        result[f"cmp_{metric}_candidate"] = str(row.get("candidate", "")).strip()
        result[f"cmp_{metric}_baseline"] = str(row.get("baseline", "")).strip()
        result[f"cmp_{metric}_diff"] = str(row.get("diff", "")).strip()
    return result


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_default_style(doc: Document) -> None:
    styles = doc.styles
    for style_name in ("Normal", "Body Text"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_paragraph(doc: Document, text: str, style: Optional[str] = None):
    paragraph = doc.add_paragraph(style=style)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("`", "")
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return paragraph


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    if width >= 1:
        add_paragraph(doc, "（表格已自动转为摘要段落，便于飞书阅读。）")
        headers = [cell.strip() for cell in rows[0]]
        for row in rows[1:]:
            parts = []
            for col_idx, header in enumerate(headers):
                if col_idx >= len(row):
                    continue
                value = row[col_idx].strip()
                if value:
                    parts.append(f"{header}: {value}")
            if parts:
                add_paragraph(doc, "；".join(parts))
        return
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    tbl_width = OxmlElement("w:tblW")
    tbl_width.set(qn("w:w"), "5000")
    tbl_width.set(qn("w:type"), "pct")
    table._tbl.tblPr.append(tbl_width)
    cell_width = Inches(6.3 / max(width, 1))
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            value = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            cell.width = cell_width
            set_cell_text(cell, value)
            if row_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def is_markdown_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").strip()
    return bool(stripped) and all(ch in "-:| " for ch in stripped)


def parse_table(lines: Sequence[str]) -> List[List[str]]:
    rows: List[List[str]] = []
    for line in lines:
        if is_markdown_table_separator(line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def resolve_markdown_image(markdown_path: Path, repo_root: Path, image_ref: str) -> Optional[Path]:
    image_ref = image_ref.strip()
    if not image_ref:
        return None
    candidate = Path(image_ref)
    candidates = []
    if candidate.is_absolute():
        candidates.append(candidate)
    else:
        candidates.append((markdown_path.parent / candidate).resolve())
        candidates.append((repo_root / candidate).resolve())
    for path in candidates:
        if path.exists() and path.is_file():
            return path
    return None


def add_markdown_image(doc: Document, image_path: Path, alt_text: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.3))
    if alt_text:
        caption = doc.add_paragraph(alt_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_markdown_to_doc(
    doc: Document,
    markdown_path: Path,
    repo_root: Optional[Path] = None,
    embed_markdown_images: bool = True,
) -> None:
    repo_root = (repo_root or Path.cwd()).resolve()
    text = markdown_path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            title = heading_match.group(2).replace("`", "")
            if level == 1:
                paragraph = doc.add_heading(title, level=0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title, level=level - 1)
            idx += 1
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if image_match:
            alt_text = image_match.group(1).strip()
            if embed_markdown_images:
                image_path = resolve_markdown_image(markdown_path, repo_root, image_match.group(2))
                if image_path is not None:
                    add_markdown_image(doc, image_path, alt_text)
                else:
                    add_paragraph(doc, f"[图片缺失] {alt_text}: {image_match.group(2)}")
            else:
                add_paragraph(doc, f"[图表位置] {alt_text}")
            idx += 1
            continue

        if line.startswith("- "):
            add_paragraph(doc, line[2:], style="List Bullet")
            idx += 1
            continue

        add_paragraph(doc, line)
        idx += 1


def summarize_nav(nav_df: Optional["pd.DataFrame"]) -> Dict[str, str]:
    if nav_df is None or nav_df.empty:
        return {}
    result: Dict[str, str] = {}
    for col in [
        "s1_active_sell_products",
        "s1_active_sell_contracts",
        "s1_active_sell_lots",
        "s1_short_open_premium_pct",
        "s1_short_liability_pct",
        "s1_put_call_lot_ratio",
        "s1_call_lot_share",
    ]:
        if col in nav_df.columns:
            result[f"{col}_mean"] = fmt_num(nav_df[col].mean(), 2)
            result[f"{col}_max"] = fmt_num(nav_df[col].max(), 2)
    if "s1_put_call_lot_ratio" in nav_df.columns:
        ratio = nav_df["s1_put_call_lot_ratio"].replace([math.inf, -math.inf], pd.NA).dropna()
        if len(ratio):
            result["pc_gt_2_days"] = str(int((ratio > 2).sum()))
            result["pc_gt_5_days"] = str(int((ratio > 5).sum()))
    return result


def get_worst_month(analysis_dir: Path) -> str:
    monthly = safe_read_csv(analysis_dir / "monthly_returns.csv")
    if monthly is None or monthly.empty:
        return "N/A"
    ret_col = "return" if "return" in monthly.columns else monthly.columns[-1]
    date_col = "month" if "month" in monthly.columns else monthly.columns[0]
    worst = monthly.sort_values(ret_col).iloc[0]
    return f"{worst[date_col]} ({fmt_pct(worst[ret_col])})"


def get_top_products(analysis_dir: Path) -> str:
    products = safe_read_csv(analysis_dir / "open_sell_product_summary.csv")
    if products is None or products.empty:
        return "N/A"
    lots_col = "open_sell_lots" if "open_sell_lots" in products.columns else products.columns[-1]
    name_col = "product" if "product" in products.columns else products.columns[0]
    top = products.sort_values(lots_col, ascending=False).head(5)
    return "、".join(f"{row[name_col]}({fmt_num(row[lots_col], 0)}手)" for _, row in top.iterrows())


def get_stop_products(analysis_dir: Path) -> str:
    stops = safe_read_csv(analysis_dir / "stop_loss_product_summary.csv")
    if stops is None or stops.empty:
        return "N/A"
    pnl_col = "net_pnl" if "net_pnl" in stops.columns else stops.columns[-1]
    name_col = "product" if "product" in stops.columns else stops.columns[0]
    top = stops.sort_values(pnl_col).head(5)
    return "、".join(f"{row[name_col]}({fmt_num(row[pnl_col], 0)})" for _, row in top.iterrows())


def get_core_close_reason(analysis_dir: Path, tag: str) -> str:
    audit_path = analysis_dir / "core_audit" / f"core_audit_{tag}.csv"
    audit = safe_read_csv(audit_path)
    if audit is None or audit.empty or "close_reason" not in audit.columns:
        return "N/A"
    audit = audit.dropna(subset=["close_reason"])
    audit = audit[audit["close_reason"].astype(str).str.strip() != ""]
    if audit.empty:
        return "N/A"
    grouped = audit.groupby("close_reason", dropna=False)["net_pnl"].agg(["count", "sum"]).sort_values("sum")
    parts = []
    for reason, row in grouped.iterrows():
        parts.append(f"{reason}: {int(row['count'])}笔, PnL {fmt_num(row['sum'], 0)}")
    return "；".join(parts)


def get_tail_product_side(analysis_dir: Path) -> str:
    data = safe_read_csv(analysis_dir / "tail_product_side_contribution.csv")
    if data is None or data.empty or "realized_pnl" not in data.columns:
        return "N/A"
    data = data.sort_values("realized_pnl").head(5)
    parts = []
    for _, row in data.iterrows():
        parts.append(f"{row.get('product', '')}-{row.get('side', '')}: {fmt_num(row.get('realized_pnl'), 0)}")
    return "；".join(parts)


def get_vega_bucket_summary(analysis_dir: Path) -> str:
    data = safe_read_csv(analysis_dir / "vega_quality_by_bucket.csv")
    if data is None or data.empty or "retained_ratio" not in data.columns:
        return "N/A"
    worst = data.sort_values("retained_ratio").head(3)
    best = data.sort_values("retained_ratio", ascending=False).head(3)
    worst_text = "、".join(
        f"{row.get('bucket', '')}-{row.get('side', '')}({fmt_pct(row.get('retained_ratio'))})"
        for _, row in worst.iterrows()
    )
    best_text = "、".join(
        f"{row.get('bucket', '')}-{row.get('side', '')}({fmt_pct(row.get('retained_ratio'))})"
        for _, row in best.iterrows()
    )
    return f"留存率较差: {worst_text}；留存率较好: {best_text}"


def get_stop_slippage_summary(analysis_dir: Path) -> str:
    stops = safe_read_csv(analysis_dir / "stop_slippage_distribution.csv")
    products = safe_read_csv(analysis_dir / "stop_slippage_product_summary.csv")
    if stops is None or stops.empty or "execution_slippage_cash" not in stops.columns:
        return "N/A"
    median_slip = stops["execution_slippage_cash"].median()
    pct_col = "close_slippage_pct_open_premium"
    median_pct = stops[pct_col].median() if pct_col in stops.columns else None
    top_product = "N/A"
    if products is not None and not products.empty and "close_slippage_cash" in products.columns:
        first = products.sort_values("close_slippage_cash", ascending=False).iloc[0]
        name = first.get("product", products.index[0] if len(products.index) else "")
        top_product = f"{name}({fmt_num(first.get('close_slippage_cash'), 0)})"
    return f"止损滑点中位数 {fmt_num(median_slip, 2)}，滑点/原权利金中位数 {fmt_pct(median_pct)}，滑点最高产品 {top_product}"


def get_pc_funnel_summary(analysis_dir: Path) -> str:
    funnel = safe_read_csv(analysis_dir / "pc_funnel.csv")
    if funnel is None or funnel.empty:
        return "N/A"
    parts = []
    for _, row in funnel.iterrows():
        parts.append(
            f"{row.get('side', '')}: 开仓权利金 {fmt_num(row.get('open_gross_premium'), 0)}, "
            f"留存率 {fmt_pct(row.get('retained_ratio'))}, 止损占比 {fmt_pct(row.get('stop_order_share'))}"
        )
    return "；".join(parts)


def build_context(tag: str, analysis_dir: Path, repo_root: Path) -> Dict[str, str]:
    metrics = read_metrics_csv(analysis_dir / "summary_metrics.csv")
    nav_path = repo_root / "output" / f"nav_{tag}.csv"
    nav = safe_read_csv(nav_path)
    nav_stats = summarize_nav(nav)
    context = {
        "total_return": fmt_pct(metrics.get("total_return")),
        "cagr": fmt_pct(metrics.get("cagr")),
        "max_drawdown": fmt_pct(metrics.get("max_drawdown")),
        "ann_vol": fmt_pct(metrics.get("ann_vol")),
        "raw_sharpe": fmt_num(metrics.get("sharpe"), 2),
        "avg_margin": fmt_pct(metrics.get("avg_margin_pct")),
        "max_margin": fmt_pct(metrics.get("max_margin_pct")),
        "cum_theta": fmt_num(metrics.get("cum_theta_pnl"), 0),
        "cum_vega": fmt_num(metrics.get("cum_vega_pnl"), 0),
        "cum_gamma": fmt_num(metrics.get("cum_gamma_pnl"), 0),
        "cum_delta": fmt_num(metrics.get("cum_delta_pnl"), 0),
        "cum_residual": fmt_num(metrics.get("cum_residual_pnl"), 0),
        "total_open_gross_premium": fmt_num(metrics.get("total_open_gross_premium"), 0),
        "avg_daily_open_gross_premium_pct_nav": fmt_pct(metrics.get("avg_daily_open_gross_premium_pct_nav")),
        "vega_pnl_to_gross_premium": fmt_pct(metrics.get("vega_pnl_to_gross_premium")),
        "vega_loss_to_gross_premium": fmt_pct(metrics.get("vega_loss_to_gross_premium")),
        "closed_premium_retained_ratio": fmt_pct(metrics.get("closed_premium_retained_ratio")),
        "worst_day": fmt_pct(metrics.get("worst_day_return")),
        "worst_month": get_worst_month(analysis_dir),
        "top_products": get_top_products(analysis_dir),
        "stop_products": get_stop_products(analysis_dir),
        "close_reason": get_core_close_reason(analysis_dir, tag),
        "tail_product_side": get_tail_product_side(analysis_dir),
        "vega_bucket_summary": get_vega_bucket_summary(analysis_dir),
        "stop_slippage_summary": get_stop_slippage_summary(analysis_dir),
        "pc_funnel_summary": get_pc_funnel_summary(analysis_dir),
    }
    context.update(nav_stats)
    context.update(read_comparison_summary(analysis_dir / "comparison_summary.csv"))
    return context


def chart_observation(chart_file: str, ctx: Dict[str, str]) -> str:
    if chart_file == "01_nav_drawdown.png":
        return (
            f"本次累计收益 {ctx.get('total_return')}，年化 {ctx.get('cagr')}，"
            f"最大回撤 {ctx.get('max_drawdown')}，最差单日 {ctx.get('worst_day')}。"
            "收益存在但斜率偏低，回撤相对目标偏深。"
        )
    if chart_file == "02_margin_positions.png":
        return (
            f"平均保证金使用率 {ctx.get('avg_margin')}，峰值 {ctx.get('max_margin')}；"
            f"平均活跃品种 {ctx.get('s1_active_sell_products_mean', 'N/A')}，"
            f"平均活跃合约 {ctx.get('s1_active_sell_contracts_mean', 'N/A')}。"
            "这说明 B0 并不是明显没用仓位，而是单位仓位产出不足。"
        )
    if chart_file in ("03_greeks_timeseries.png", "04_pnl_attribution.png"):
        return (
            f"Theta 累计 {ctx.get('cum_theta')}，Vega 累计 {ctx.get('cum_vega')}，"
            f"Gamma 累计 {ctx.get('cum_gamma')}，Delta 累计 {ctx.get('cum_delta')}，"
            f"Residual 累计 {ctx.get('cum_residual')}。"
            "卖方 carry 很明确，但 vega/gamma 消耗太大。"
        )
    if chart_file == "05_daily_pnl_tail.png":
        return f"最差单日收益 {ctx.get('worst_day')}，最差月份为 {ctx.get('worst_month')}。需要检查这些日期是否对应升波、跳空和止损簇集。"
    if chart_file == "06_premium_pc_structure.png":
        return (
            f"P/C 比均值 {ctx.get('s1_put_call_lot_ratio_mean', 'N/A')}，峰值 {ctx.get('s1_put_call_lot_ratio_max', 'N/A')}；"
            f"P/C>2 的天数 {ctx.get('pc_gt_2_days', 'N/A')}，P/C>5 的天数 {ctx.get('pc_gt_5_days', 'N/A')}。"
            "这提示 B0 虽然定义为双侧卖权，但实际暴露会阶段性单侧化。"
        )
    if chart_file == "08_calendar_returns.png":
        return f"月度维度最差月份为 {ctx.get('worst_month')}。这类月份应作为后续规则优化的压力样本，而不是被平均收益掩盖。"
    if chart_file == "09_product_share_top10.png":
        return f"开仓手数最高的品种为 {ctx.get('top_products')}。如果这些品种集中在同一板块，需要用板块和相关性预算约束。"
    if chart_file in ("10_order_action_summary.png", "11_close_event_timeline.png"):
        return f"平仓路径摘要：{ctx.get('close_reason')}。止损亏损集中品种包括 {ctx.get('stop_products')}。"
    if chart_file == "12_daily_open_premium_vega_quality.png":
        return (
            f"新开仓毛权利金合计 {ctx.get('total_open_gross_premium')}，"
            f"日均新开仓毛权利金/NAV {ctx.get('avg_daily_open_gross_premium_pct_nav')}，"
            f"Vega PnL/毛权利金 {ctx.get('vega_pnl_to_gross_premium')}，"
            f"Vega loss/毛权利金 {ctx.get('vega_loss_to_gross_premium')}，"
            f"已平仓权利金留存率 {ctx.get('closed_premium_retained_ratio')}。"
            "如果这组指标显示 vega 吞噬过高，后续优化应优先控制 vega，而不是简单提高开仓权利金流量。"
        )
    if chart_file == "13_tail_product_side_contribution.png":
        return (
            f"尾部日平仓实现亏损最集中的产品/方向为：{ctx.get('tail_product_side')}。"
            "这张图只能解释尾部日的实现平仓损益，若与 NAV 最差日不完全闭合，应继续做产品级盯市归因。"
        )
    if chart_file == "14_vega_quality_by_bucket.png":
        return (
            f"分桶权利金留存摘要：{ctx.get('vega_bucket_summary')}。"
            "若高波或升波桶留存率差，说明权利金厚度没有补偿 vega 风险，应该降低该桶预算。"
        )
    if chart_file == "15_stop_slippage_distribution.png":
        return (
            f"止损滑点摘要：{ctx.get('stop_slippage_summary')}。"
            "滑点分布越厚，越说明回测里的止损执行需要保守化，并对低流动性产品加额外惩罚。"
        )
    if chart_file == "16_pc_funnel.png":
        return (
            f"Put/Call 漏斗摘要：{ctx.get('pc_funnel_summary')}。"
            "如果某一侧开仓权利金占比高但留存率低，说明 P/C 偏移可能只是增加了承保量，而没有提高承保质量。"
        )
    if chart_file == "compare_01_nav_relative_to_b0.png":
        return (
            f"相对 B0，累计收益差 {fmt_pct(ctx.get('cmp_total_return_diff'))}，"
            f"年化差 {fmt_pct(ctx.get('cmp_cagr_diff'))}，Sharpe 差 {fmt_num(ctx.get('cmp_sharpe_diff'), 2)}。"
            "若超额收益曲线不连续，需要回到订单和品种贡献里确认是否来自少数路径。"
        )
    if chart_file == "compare_02_drawdown_relative_to_b0.png":
        return (
            f"候选版本相对 B0 的最大回撤差 {fmt_pct(ctx.get('cmp_max_drawdown_diff'))}，"
            f"最差单日收益差 {fmt_pct(ctx.get('cmp_worst_day_return_diff'))}。"
            "负值代表候选版本左尾更深，正值代表回撤更浅。"
        )
    if chart_file == "compare_03_margin_position_relative_to_b0.png":
        return (
            f"平均保证金使用率差 {fmt_pct(ctx.get('cmp_avg_margin_pct_diff'))}，"
            f"峰值保证金使用率差 {fmt_pct(ctx.get('cmp_max_margin_pct_diff'))}。"
            "这决定了收益改善是来自风险预算提高，还是来自同等仓位下的选约效率提升。"
        )
    if chart_file == "compare_04_greek_attribution_relative_to_b0.png":
        return (
            f"相对 B0，theta 差 {fmt_num(ctx.get('cmp_cum_theta_pnl_diff'), 0)}，"
            f"vega 差 {fmt_num(ctx.get('cmp_cum_vega_pnl_diff'), 0)}，"
            f"gamma 差 {fmt_num(ctx.get('cmp_cum_gamma_pnl_diff'), 0)}，"
            f"delta 差 {fmt_num(ctx.get('cmp_cum_delta_pnl_diff'), 0)}。"
            "这是判断候选版本是否真的更像卖波策略的第一优先证据。"
        )
    if chart_file == "compare_05_daily_pnl_tail_relative_to_b0.png":
        return (
            f"最差单日差 {fmt_pct(ctx.get('cmp_worst_day_return_diff'))}，"
            "需要结合分布图确认候选版本是否把左尾推厚。若左尾恶化，收益提升不能直接外推。"
        )
    if chart_file == "compare_06_pc_structure_relative_to_b0.png":
        return (
            f"P/C 均值差 {fmt_num(ctx.get('cmp_avg_s1_put_call_lot_ratio_diff'), 2)}，"
            f"Call 手数占比均值差 {fmt_num(ctx.get('cmp_avg_s1_call_lot_share_diff'), 2)}。"
            "若候选版本 P/C 偏离显著扩大，应进一步拆 Put 和 Call 的独立贡献。"
        )
    if chart_file == "compare_07_product_exposure_relative_to_b0.png":
        return (
            "这张图应重点看 Top 产品是否更集中，以及产品贡献差异是否由少数品种主导。"
            "若集中度上升，需要在下一轮报告中同步检查板块和相关组暴露。"
        )
    if chart_file == "compare_08_stop_cluster_relative_to_b0.png":
        return (
            "这张图用于确认候选版本是否减少止损簇集，或只是把止损延后。"
            "若止损金额下降但最差日不改善，说明尾部风险可能从显性止损转为持仓浮亏。"
        )
    return "该图用于补充观察策略结构与收益路径，具体结论应结合主文指标和交易流水共同判断。"


def add_chart_section(doc: Document, tag: str, analysis_dir: Path, repo_root: Path) -> None:
    context = build_context(tag, analysis_dir, repo_root)
    doc.add_page_break()
    doc.add_heading("图表展示附录（机器快读，不替代正文深度分析）", level=1)
    add_paragraph(
        doc,
        "本节只负责把分析包中的标准 PNG 稳定嵌入 DOCX，并给出便于检索的机器快读。"
        "它不能替代报告正文里的逐图深度分析。正式汇报必须在正文中逐张解释图表："
        "读图口径、图上证据、期权策略含义、可能的回测口径问题和下一步验证。",
    )

    for idx, spec in enumerate(CHART_SPECS, start=1):
        image_path = analysis_dir / spec["file"]
        if not image_path.exists():
            continue
        doc.add_heading(f"图 {idx}：{spec['title']}", level=2)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.3))
        caption = doc.add_paragraph(f"图 {idx}：{spec['title']}（来源：{image_path.name}）")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph(doc, f"机器快读 - 读图口径：{spec['read']}")
        add_paragraph(doc, f"机器快读 - 本次观察：{chart_observation(spec['file'], context)}")
        add_paragraph(doc, f"机器快读 - 策略提示：{spec['meaning']}")


def add_metadata(doc: Document, title: str, tag: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"版本：{tag}")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = title
    footer = section.footer.paragraphs[0]
    footer.text = "S1 backtest attribution report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_table_borders(doc: Document) -> None:
    for table in doc.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "D9E2F3")
            borders.append(border)
        tbl_pr.append(borders)


def main() -> None:
    args = parse_args()
    repo_root = args.repo_root.resolve()
    analysis_dir = (args.analysis_dir or (repo_root / "output" / f"analysis_{args.tag}")).resolve()
    output_docx = args.output_docx or (analysis_dir / f"{args.tag}_feishu_report.docx")
    title = args.title or "S1 回测归因与策略复盘"

    if not args.markdown.exists():
        raise SystemExit(f"Markdown report not found: {args.markdown}")
    if not analysis_dir.exists():
        raise SystemExit(f"Analysis directory not found: {analysis_dir}")

    doc = Document()
    set_default_style(doc)
    add_markdown_to_doc(
        doc,
        args.markdown,
        repo_root=repo_root,
        embed_markdown_images=not args.skip_markdown_images,
    )
    add_metadata(doc, title, args.tag)
    if not args.skip_existing_chart_section:
        add_chart_section(doc, args.tag, analysis_dir, repo_root)
    set_table_borders(doc)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    print(str(output_docx))


if __name__ == "__main__":
    main()
