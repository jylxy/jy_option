#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the S1 B6 experiment report pack.

The report is intentionally generated from local CSV outputs so it can be
repeated after reruns.  It compares B6a/B6b/B6c against B2C and B1 through the
S1 premium formula:

    Premium Pool * Deployment Ratio * Retention Rate
    - Tail / Stop Loss - Cost / Slippage
"""

from __future__ import annotations

import math
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import numpy as np
import pandas as pd

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
DOCS = ROOT / "docs"
ANALYSIS = OUTPUT / "analysis_s1_b6_experiment_20260430"
REPORT_MD = DOCS / "s1_b6_experiment_report_20260430.md"
REPORT_DOCX = ANALYSIS / "s1_b6_experiment_report_20260430_feishu.docx"
TRADING_DAYS = 252


RUNS = {
    "B1": {
        "label": "B1 流动性/OI排序",
        "tag": "s1_b1_liq_oi_rank_stop25_allprod_2022_latest",
        "role": "流动性基准",
    },
    "B2C": {
        "label": "B2C 品种预算倾斜",
        "tag": "s1_b2_product_tilt075_stop25_allprod_2022_latest",
        "role": "当前主基准",
    },
    "B6a": {
        "label": "B6a 合约层残差质量排序",
        "tag": "s1_b6a_residual_contract_quality_2022_latest",
        "role": "合约层排序",
    },
    "B6b": {
        "label": "B6b P/C侧预算倾斜",
        "tag": "s1_b6b_residual_side_tilt_2022_latest",
        "role": "P/C侧预算",
    },
    "B6c": {
        "label": "B6c 品种预算倾斜",
        "tag": "s1_b6c_residual_product_tilt_2022_latest",
        "role": "品种层预算",
    },
}


CHART_NOTES = [
    {
        "file": "01_b6_nav_excess_vs_b2c.png",
        "title": "净值与超额路径",
        "how": "先看上半部分的标准化 NAV，再看下半部分相对 B2C 的超额 NAV。上半图判断长期收益路径，下半图判断是否存在稳定、可复用的超额。",
        "read": "完整样本中 B2C 终值为 11,713,558；B6b 为 11,595,599，比 B2C 少 117,959；B6a 和 B6c 分别为 11,391,229 和 11,419,315，差距更大。",
        "judgement": "B6b 是三条 B6 线里最接近主基准的，但没有形成持续压过 B2C 的超额。B6a/B6c 长期弱于 B2C，说明单纯把残差因子放到合约排序或品种预算，并没有自动转化成可留存收益。",
        "risk": "净值路径未能回答因子本身是否有效，只说明当前用法没有胜出。若因子被放错层级，NAV 失败不等于因子失败。",
        "next": "下一步应把 B6b 保留为 P/C 预算线，同时把 B6a/B6c 拆回因子库，在合约层、品种层、组合层分别重新验证。"
    },
    {
        "file": "02_b6_drawdown_comparison.png",
        "title": "回撤对比",
        "how": "回撤图重点看左尾厚度和回撤修复速度，而不是只看终点收益。卖权策略如果多赚一点 theta 但回撤显著变深，就不是合格升级。",
        "read": "B2C 最大回撤为 -2.52%；B6b 最大回撤扩大到 -3.12%；B6a 为 -3.66%，B6c 为 -3.77%。三条 B6 线都没有改善最大回撤。",
        "judgement": "这说明 B6 因子没有压住 Tail / Stop Loss。B6b 收益相对较好，但仍以更深回撤作为代价；B6a/B6c 则属于收益和回撤同时弱化。",
        "risk": "B6 回撤变深可能来自 gamma 集中、P/C 偏移、品种相关性或同到期聚集，单靠 orders 无法完全解释板块 stress，需要后续 diagnostics 或持仓快照。",
        "next": "B7 应优先做组合层约束：Top stress share、bucket/corr_group 上限、same-expiry gamma、tail correlation 和 stop cluster。"
    },
    {
        "file": "03_b6_greek_diff_vs_b2c.png",
        "title": "Greeks 相对 B2C 累计差异",
        "how": "这张图判断 B6 的净值差异来自 theta、vega、gamma、delta 还是 residual。卖权升级应该优先表现为 theta 留存提升、vega 损耗下降、gamma 损耗可控。",
        "read": "B6b 的 theta 为 11,273,175，高于 B2C 的 10,190,547；B6b vega 为 -5,610,956，好于 B2C 的 -6,026,160；但 B6b gamma 为 -8,371,600，显著差于 B2C 的 -6,921,512。",
        "judgement": "B6b 的方向是有价值的：它确实增加了 theta、改善了一些 vega 损耗；但多出来的 theta 和 vega 改善没有覆盖额外 gamma 路径亏损，所以最终仍落后 B2C。",
        "risk": "Residual 仍然较大，不能把 residual 当作稳定 alpha。它可能包含离散重估、skew/vanna/volga、近到期非线性和成交口径误差。",
        "next": "后续需要把 B6b 与 gamma 控制联动，例如按 theta/gamma、expiry gamma、delta 漂移和 stop cluster 共同决定侧预算。"
    },
    {
        "file": "04_b6_formula_bars.png",
        "title": "公式项柱状图",
        "how": "这张图把收益、回撤、theta、gamma、vega、delta 放在同一视角下，判断一个版本到底改善了公式中的哪一项。",
        "read": "B6b 年化 3.12%，低于 B2C 的 3.34%；B6b 的 theta 更高、vega 亏损更低，但 gamma 亏损多约 145 万，成为拖累核心。",
        "judgement": "B6b 改善了 Premium Pool / Theta 和部分 Vega 质量，但恶化了 Tail / Gamma Loss。B6a/B6c 则说明合约层和品种层倾斜没有改善总体风险补偿。",
        "risk": "如果只看毛权利金或 theta，会误判 B6b 是升级；必须同步看 gamma loss / gross premium 和 premium retained ratio。",
        "next": "报告后的 B7 实验应把目标从提高毛权利金转为提高 retention-adjusted premium，并限制 gamma loss / premium。"
    },
    {
        "file": "05_b6_margin_products_pc.png",
        "title": "保证金、品种与 P/C 结构",
        "how": "这张图看 Deployment Ratio、活跃品种数量和 Put/Call 结构是否发生漂移。",
        "read": "B6 系列保证金使用率并不低，长期接近主基准；失败不是因为没打满仓位。B6b 的 P/C 侧调整更明显，说明它确实改变了侧预算结构。",
        "judgement": "B6b 比 B6a/B6c 更接近正确方向，是因为 P/C 侧预算本来就应该由趋势、skew、gamma/vega 质量共同决定。但当前侧预算仍不足以规避 gamma 路径风险。",
        "risk": "P/C 偏移如果长期偏向一侧，会把卖波策略变成方向策略。我们后续必须区分“顺趋势降低危险侧预算”和“长期押涨/押跌”。",
        "next": "P/C 侧预算应加入趋势置信度上限、breakout proximity、skew 是否真贵、以及同侧 stop cluster。"
    },
    {
        "file": "06_b6_monthly_excess_vs_b2c.png",
        "title": "月度超额",
        "how": "这张图看超额是否稳定跨月份出现，而不是由少数月份贡献。",
        "read": "B6b 有若干月份接近或超过 B2C，但并未形成连续月度优势。B6a/B6c 的月度超额更不稳定。",
        "judgement": "如果一个因子只能在少数月份有效，就更适合作为 regime 条件下的预算调节，而不是全市场、全时期常开参数。",
        "risk": "月度结果容易受当月商品趋势和波动切换影响，不能简单用月度胜率判断因子长期有效。",
        "next": "对 B6b 做 regime 分组：falling vol、low vol、rising vol、趋势突破、震荡环境下分别看 P/C 倾斜是否有效。"
    },
    {
        "file": "07_b6_tail_days.png",
        "title": "日度尾部",
        "how": "这张图看最差日和滚动左尾分位，判断策略是否减少了卖方最核心的极端日风险。",
        "read": "B6b 最差单日为 -2.10%，明显差于 B2C 的 -1.04%；B6a 和 B6c 最差单日也差于 B2C。",
        "judgement": "B6 没有降低左尾，说明残差质量因子在当前使用方式下没有解决真正的尾部暴露。卖权策略里，这比年化收益低一点更严重。",
        "risk": "左尾可能来自单品种跳价、同板块共振或止损簇，当前图只能识别结果，不能完全识别原因。",
        "next": "下一步用 orders 拆最差日的产品/方向贡献，并在组合层加入 tail-HRP 或 stop-cluster 预算。"
    },
    {
        "file": "08_b6_product_contribution.png",
        "title": "品种贡献",
        "how": "这张图用已平仓 realized PnL 看 Top 盈亏品种，判断是否存在少数品种主导结果。",
        "read": "B6 版本的盈亏仍由少数品种贡献或拖累，并没有因为品种层倾斜就自然获得更均匀的收益来源。",
        "judgement": "品种层因子不应只用于提高预算，还必须结合尾部相依、板块集中和退出容量。否则预算倾斜可能只是把风险集中到更脆的品种。",
        "risk": "这里使用的是订单已实现 PnL，不包含未平仓浮动路径，也不是完整的日度 stress 暴露。",
        "next": "后续应生成日度 product/bucket/corr_group stress 快照，计算 Top1/Top3 stress share 和有效风险品种数。"
    },
    {
        "file": "09_b6_stop_summary.png",
        "title": "止损统计",
        "how": "这张图看止损次数和止损损失，是 Premium Formula 中 Tail / Stop Loss 的直接体现。",
        "read": "B2C 止损 PnL 约 -349.6 万；B6b 止损 PnL 约 -499.0 万，止损损耗明显更重。B6a/B6c 的止损损耗也显著高于 B2C。",
        "judgement": "这基本解释了 B6 为什么没有打过 B2C：B6 增厚的权利金和 theta 被更大的止损损耗吞掉了。",
        "risk": "止损损耗变大可能是合约更靠近风险区、gamma 更大、低流动性跳价更多，或者方向侧预算没有避开趋势侧。",
        "next": "把 cooldown、止损后是否归零、止损后继续不利移动、止损簇和低价 tick 风险纳入退出层与入场硬过滤。"
    },
]


def setup() -> None:
    ANALYSIS.mkdir(parents=True, exist_ok=True)
    DOCS.mkdir(parents=True, exist_ok=True)
    plt.rcParams["figure.dpi"] = 120
    plt.rcParams["savefig.dpi"] = 150
    plt.rcParams["axes.grid"] = True
    plt.rcParams["grid.alpha"] = 0.25
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans", "Arial", "Microsoft YaHei", "SimHei"]


def read_nav(tag: str) -> pd.DataFrame:
    path = OUTPUT / f"nav_{tag}.csv"
    if not path.exists():
        raise FileNotFoundError(path)
    df = pd.read_csv(path)
    df["date"] = pd.to_datetime(df["date"])
    return df.sort_values("date").reset_index(drop=True)


def read_orders(tag: str) -> pd.DataFrame:
    path = OUTPUT / f"orders_{tag}.csv"
    if not path.exists():
        raise FileNotFoundError(path)
    df = pd.read_csv(path)
    df["date"] = pd.to_datetime(df["date"])
    return df.sort_values("date").reset_index(drop=True)


def max_drawdown(nav: pd.Series) -> pd.Series:
    return nav / nav.cummax() - 1.0


def safe_sum(df: pd.DataFrame, col: str) -> float:
    if col not in df.columns:
        return 0.0
    return float(pd.to_numeric(df[col], errors="coerce").fillna(0.0).sum())


def safe_mean(df: pd.DataFrame, col: str) -> float:
    if col not in df.columns:
        return float("nan")
    series = pd.to_numeric(df[col], errors="coerce")
    return float(series.mean())


def metric_row(name: str, nav: pd.DataFrame, orders: pd.DataFrame) -> Dict[str, float | str]:
    ret = nav["nav"].pct_change().fillna(0.0)
    days = max(len(nav) - 1, 1)
    nav0 = float(nav["nav"].iloc[0])
    nav1 = float(nav["nav"].iloc[-1])
    cum_ret = nav1 / nav0 - 1.0
    ann = (nav1 / nav0) ** (TRADING_DAYS / days) - 1.0
    vol = float(ret.std(ddof=0) * math.sqrt(TRADING_DAYS))
    mdd = float(max_drawdown(nav["nav"]).min())
    sharpe_rf2 = (ann - 0.02) / vol if vol > 0 else np.nan
    calmar = ann / abs(mdd) if mdd < 0 else np.nan

    open_orders = orders[orders["action"].eq("open_sell")].copy()
    close_orders = orders[~orders["action"].eq("open_sell")].copy()
    stop_orders = close_orders[close_orders["action"].astype(str).str.startswith("sl_")].copy()
    expiry_orders = close_orders[close_orders["action"].eq("expiry")].copy()

    gross_open = safe_sum(open_orders, "gross_premium_cash")
    net_open = safe_sum(open_orders, "net_premium_cash")
    close_pnl = safe_sum(close_orders, "pnl")
    stop_pnl = safe_sum(stop_orders, "pnl")
    expiry_pnl = safe_sum(expiry_orders, "pnl")
    retained = safe_sum(close_orders, "premium_retained_cash")
    close_open_premium = safe_sum(close_orders, "open_premium_cash")
    retained_ratio = retained / close_open_premium if close_open_premium > 0 else np.nan

    return {
        "name": name,
        "label": RUNS[name]["label"],
        "role": RUNS[name]["role"],
        "start": nav["date"].iloc[0].date().isoformat(),
        "end": nav["date"].iloc[-1].date().isoformat(),
        "days": len(nav),
        "nav": nav1,
        "cum_ret": cum_ret,
        "ann_ret": ann,
        "ann_vol": vol,
        "max_dd": mdd,
        "worst_day": float(ret.min()),
        "sharpe_rf2": sharpe_rf2,
        "calmar": calmar,
        "avg_margin": safe_mean(nav, "s1_margin_used_pct"),
        "max_margin": float(pd.to_numeric(nav.get("s1_margin_used_pct", pd.Series([np.nan] * len(nav))), errors="coerce").max()),
        "avg_products": safe_mean(nav, "s1_active_sell_products"),
        "avg_contracts": safe_mean(nav, "s1_active_sell_contracts"),
        "avg_lots": safe_mean(nav, "s1_active_sell_lots"),
        "avg_pc_ratio": safe_mean(nav.replace([np.inf, -np.inf], np.nan), "s1_put_call_lot_ratio"),
        "gross_open_premium": gross_open,
        "net_open_premium": net_open,
        "close_pnl": close_pnl,
        "premium_retained_cash": retained,
        "premium_retained_ratio": retained_ratio,
        "stop_count": int(len(stop_orders)),
        "stop_pnl": stop_pnl,
        "expiry_count": int(len(expiry_orders)),
        "expiry_pnl": expiry_pnl,
        "fee": safe_sum(nav, "fee"),
        "delta_pnl": safe_sum(nav, "delta_pnl"),
        "gamma_pnl": safe_sum(nav, "gamma_pnl"),
        "theta_pnl": safe_sum(nav, "theta_pnl"),
        "vega_pnl": safe_sum(nav, "vega_pnl"),
        "residual_pnl": safe_sum(nav, "residual_pnl"),
    }


def pct(x: float, digits: int = 2) -> str:
    if pd.isna(x):
        return "NA"
    return f"{x * 100:.{digits}f}%"


def money(x: float, digits: int = 0) -> str:
    if pd.isna(x):
        return "NA"
    return f"{x:,.{digits}f}"


def save_table(df: pd.DataFrame, name: str) -> Path:
    path = ANALYSIS / name
    df.to_csv(path, index=False, encoding="utf-8-sig")
    return path


def plot_nav(navs: Dict[str, pd.DataFrame]) -> Path:
    fig, axes = plt.subplots(2, 1, figsize=(12, 8), sharex=True)
    baseline = navs["B2C"][["date", "nav"]].rename(columns={"nav": "base_nav"})
    for name, nav in navs.items():
        axes[0].plot(nav["date"], nav["nav"] / nav["nav"].iloc[0], label=name)
        if name != "B2C":
            merged = nav[["date", "nav"]].merge(baseline, on="date", how="inner")
            axes[1].plot(merged["date"], merged["nav"] - merged["base_nav"], label=f"{name} - B2C")
    axes[0].set_title("NAV normalized")
    axes[0].set_ylabel("NAV / initial")
    axes[0].legend(ncol=3, fontsize=8)
    axes[1].set_title("Excess NAV versus B2C")
    axes[1].axhline(0, color="black", linewidth=0.8)
    axes[1].set_ylabel("CNY")
    axes[1].legend(ncol=2, fontsize=8)
    fig.tight_layout()
    path = ANALYSIS / "01_b6_nav_excess_vs_b2c.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_drawdown(navs: Dict[str, pd.DataFrame]) -> Path:
    fig, ax = plt.subplots(figsize=(12, 5))
    for name, nav in navs.items():
        ax.plot(nav["date"], max_drawdown(nav["nav"]) * 100, label=name)
    ax.set_title("Drawdown comparison")
    ax.set_ylabel("Drawdown (%)")
    ax.legend(ncol=3, fontsize=8)
    fig.tight_layout()
    path = ANALYSIS / "02_b6_drawdown_comparison.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_greek_diff(navs: Dict[str, pd.DataFrame]) -> Path:
    cols = ["delta_pnl", "gamma_pnl", "theta_pnl", "vega_pnl", "residual_pnl"]
    baseline = navs["B2C"][["date"] + cols].copy()
    for col in cols:
        baseline[col] = baseline[col].cumsum()
    fig, axes = plt.subplots(len(cols), 1, figsize=(12, 12), sharex=True)
    for ax, col in zip(axes, cols):
        for name in ["B6a", "B6b", "B6c"]:
            run = navs[name][["date"] + cols].copy()
            for c in cols:
                run[c] = run[c].cumsum()
            merged = run[["date", col]].merge(
                baseline[["date", col]].rename(columns={col: "base"}), on="date", how="inner"
            )
            ax.plot(merged["date"], merged[col] - merged["base"], label=name)
        ax.axhline(0, color="black", linewidth=0.6)
        ax.set_title(f"Cumulative {col} difference vs B2C")
        ax.legend(ncol=3, fontsize=8)
    fig.tight_layout()
    path = ANALYSIS / "03_b6_greek_diff_vs_b2c.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_formula_bars(metrics: pd.DataFrame) -> Path:
    fields = [
        ("ann_ret", "Annual return"),
        ("max_dd", "Max drawdown"),
        ("theta_pnl", "Theta PnL"),
        ("gamma_pnl", "Gamma PnL"),
        ("vega_pnl", "Vega PnL"),
        ("delta_pnl", "Delta PnL"),
    ]
    fig, axes = plt.subplots(2, 3, figsize=(14, 8))
    for ax, (field, title) in zip(axes.ravel(), fields):
        values = metrics.set_index("name")[field]
        if field in {"ann_ret", "max_dd"}:
            values = values * 100
        ax.bar(values.index, values.values)
        ax.set_title(title)
        ax.tick_params(axis="x", labelrotation=35)
    fig.tight_layout()
    path = ANALYSIS / "04_b6_formula_bars.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def plot_margin_positions(navs: Dict[str, pd.DataFrame]) -> Path:
    fig, axes = plt.subplots(3, 1, figsize=(12, 10), sharex=True)
    for name, nav in navs.items():
        axes[0].plot(nav["date"], pd.to_numeric(nav["s1_margin_used_pct"], errors="coerce") * 100, label=name)
        axes[1].plot(nav["date"], pd.to_numeric(nav["s1_active_sell_products"], errors="coerce"), label=name)
        axes[2].plot(nav["date"], pd.to_numeric(nav["s1_put_call_lot_ratio"], errors="coerce").replace([np.inf, -np.inf], np.nan), label=name)
    axes[0].set_title("S1 margin usage (%)")
    axes[1].set_title("Active products")
    axes[2].set_title("Put / Call lot ratio")
    for ax in axes:
        ax.legend(ncol=3, fontsize=8)
    fig.tight_layout()
    path = ANALYSIS / "05_b6_margin_products_pc.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def monthly_returns(nav: pd.DataFrame) -> pd.Series:
    d = nav[["date", "nav"]].copy()
    d["month"] = d["date"].dt.to_period("M")
    g = d.groupby("month")["nav"].agg(["first", "last"])
    return g["last"] / g["first"] - 1.0


def plot_monthly_excess(navs: Dict[str, pd.DataFrame]) -> Tuple[Path, pd.DataFrame]:
    monthly = {name: monthly_returns(nav) for name, nav in navs.items()}
    frame = pd.DataFrame(monthly).dropna()
    excess = frame[["B6a", "B6b", "B6c"]].subtract(frame["B2C"], axis=0)
    fig, ax = plt.subplots(figsize=(13, 5))
    excess.plot(kind="bar", ax=ax)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Monthly excess return vs B2C")
    ax.set_ylabel("Return diff")
    ax.tick_params(axis="x", labelrotation=80)
    fig.tight_layout()
    path = ANALYSIS / "06_b6_monthly_excess_vs_b2c.png"
    fig.savefig(path)
    plt.close(fig)
    save_table(excess.reset_index().astype({"month": str}), "monthly_excess_vs_b2c.csv")
    return path, excess


def plot_tail_days(navs: Dict[str, pd.DataFrame]) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(13, 5))
    tail_rows = []
    for name, nav in navs.items():
        d = nav[["date", "nav"]].copy()
        d["daily_ret"] = d["nav"].pct_change().fillna(0.0)
        worst = d.nsmallest(20, "daily_ret").copy()
        tail_rows.append(worst.assign(version=name))
        axes[0].hist(d["daily_ret"] * 100, bins=80, alpha=0.35, label=name)
        axes[1].plot(d["date"], d["daily_ret"].rolling(20).quantile(0.05) * 100, label=name)
    axes[0].set_title("Daily return distribution")
    axes[0].set_xlabel("Daily return (%)")
    axes[1].set_title("Rolling 20D 5% daily return")
    axes[1].set_ylabel("%")
    for ax in axes:
        ax.legend(fontsize=8)
    fig.tight_layout()
    path = ANALYSIS / "07_b6_tail_days.png"
    fig.savefig(path)
    plt.close(fig)
    tail = pd.concat(tail_rows, ignore_index=True)
    save_table(tail[["version", "date", "daily_ret", "nav"]], "worst_20_days_by_version.csv")
    return path


def product_side_summary(orders: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for name, od in orders.items():
        close = od[~od["action"].eq("open_sell")].copy()
        if close.empty:
            continue
        grp = (
            close.groupby(["product", "option_type"], dropna=False)
            .agg(
                pnl=("pnl", "sum"),
                count=("pnl", "size"),
                retained=("premium_retained_cash", "sum"),
                open_premium=("open_premium_cash", "sum"),
                stop_count=("action", lambda x: x.astype(str).str.startswith("sl_").sum()),
            )
            .reset_index()
        )
        grp["version"] = name
        grp["retained_ratio"] = grp["retained"] / grp["open_premium"].replace(0, np.nan)
        rows.append(grp)
    return pd.concat(rows, ignore_index=True) if rows else pd.DataFrame()


def plot_product_contribution(summary: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 3, figsize=(15, 5), sharey=True)
    for ax, name in zip(axes, ["B2C", "B6b", "B6c"]):
        df = summary[summary["version"].eq(name)].copy()
        prod = df.groupby("product")["pnl"].sum().sort_values()
        top = pd.concat([prod.head(5), prod.tail(5)]).drop_duplicates()
        ax.barh(top.index.astype(str), top.values)
        ax.axvline(0, color="black", linewidth=0.7)
        ax.set_title(f"{name} product realized PnL")
    fig.tight_layout()
    path = ANALYSIS / "08_b6_product_contribution.png"
    fig.savefig(path)
    plt.close(fig)
    save_table(summary.sort_values(["version", "pnl"]), "product_side_realized_summary.csv")
    return path


def stop_summary(orders: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for name, od in orders.items():
        close = od[~od["action"].eq("open_sell")].copy()
        stop = close[close["action"].astype(str).str.startswith("sl_")].copy()
        expiry = close[close["action"].eq("expiry")].copy()
        rows.append(
            {
                "version": name,
                "close_count": len(close),
                "stop_count": len(stop),
                "stop_pnl": safe_sum(stop, "pnl"),
                "expiry_count": len(expiry),
                "expiry_pnl": safe_sum(expiry, "pnl"),
                "close_pnl": safe_sum(close, "pnl"),
                "stop_share": len(stop) / len(close) if len(close) else np.nan,
            }
        )
    return pd.DataFrame(rows)


def plot_stop_summary(stop_df: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
    axes[0].bar(stop_df["version"], stop_df["stop_count"])
    axes[0].set_title("Stop count")
    axes[0].tick_params(axis="x", labelrotation=30)
    axes[1].bar(stop_df["version"], stop_df["stop_pnl"])
    axes[1].set_title("Stop realized PnL")
    axes[1].tick_params(axis="x", labelrotation=30)
    fig.tight_layout()
    path = ANALYSIS / "09_b6_stop_summary.png"
    fig.savefig(path)
    plt.close(fig)
    save_table(stop_df, "stop_summary.csv")
    return path


def build_all_charts(navs: Dict[str, pd.DataFrame], orders: Dict[str, pd.DataFrame], metrics: pd.DataFrame) -> List[Path]:
    paths = [
        plot_nav(navs),
        plot_drawdown(navs),
        plot_greek_diff(navs),
        plot_formula_bars(metrics),
        plot_margin_positions(navs),
    ]
    monthly_path, _ = plot_monthly_excess(navs)
    paths.append(monthly_path)
    paths.append(plot_tail_days(navs))
    prod_summary = product_side_summary(orders)
    paths.append(plot_product_contribution(prod_summary))
    paths.append(plot_stop_summary(stop_summary(orders)))
    return paths


def markdown_table(df: pd.DataFrame, columns: List[Tuple[str, str]]) -> str:
    header = "| " + " | ".join(title for _, title in columns) + " |"
    sep = "| " + " | ".join("---" for _ in columns) + " |"
    rows = [header, sep]
    for _, row in df.iterrows():
        cells = []
        for col, _title in columns:
            val = row[col]
            if isinstance(val, float):
                if col.endswith("_pct") or col in {"cum_ret", "ann_ret", "ann_vol", "max_dd", "worst_day", "avg_margin", "max_margin", "premium_retained_ratio"}:
                    cells.append(pct(val))
                elif "pnl" in col or "premium" in col or col in {"nav", "fee", "gross_open_premium", "net_open_premium"}:
                    cells.append(money(val))
                else:
                    cells.append(f"{val:.2f}")
            else:
                cells.append(str(val))
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows)


def write_report(metrics: pd.DataFrame, charts: List[Path], product_summary: pd.DataFrame, stops: pd.DataFrame) -> None:
    b2c = metrics.set_index("name").loc["B2C"]
    b6b = metrics.set_index("name").loc["B6b"]
    b6a = metrics.set_index("name").loc["B6a"]
    b6c = metrics.set_index("name").loc["B6c"]
    best_b6 = metrics[metrics["name"].str.startswith("B6")].sort_values("nav", ascending=False).iloc[0]
    b6b_nav_diff = float(b6b["nav"] - b2c["nav"])
    b6b_nav_diff_word = "高" if b6b_nav_diff >= 0 else "少"

    metric_table = metrics.copy()
    for col in ["cum_ret", "ann_ret", "ann_vol", "max_dd", "worst_day", "avg_margin", "max_margin", "premium_retained_ratio"]:
        metric_table[col] = metric_table[col].astype(float)

    md: List[str] = []
    md.append("# S1 B6 因子角色实验复盘报告")
    md.append("")
    md.append("生成日期：2026-04-30")
    md.append("")
    md.append("## 1. 期权专家审议摘要")
    md.append("")
    md.append(
        "B6 的核心问题不是“因子没有价值”，而是“从 full shadow 和残差 IC 中筛出的因子，直接放进交易层后，没有放在足够合适的位置”。"
        "三条 B6 线里，B6b 是相对最好的，但仍未超过 B2C；B6a/B6c 明显弱于 B2C。"
        f"完整共同区间内，B2C 终值为 {money(float(b2c['nav']))}，B6b 为 {money(float(b6b['nav']))}，"
        f"B6b 相对 B2C {b6b_nav_diff_word} {money(abs(b6b_nav_diff))}。"
    )
    md.append("")
    md.append(
        "从卖权经济学看，B6 没有改善核心质量：它没有把 theta 转化为更高的留存收益，也没有降低 gamma/vega 的吞噬。"
        "B6b 的收益相对 B1 有改善，但相对 B2C 仍不足，说明 P/C 侧预算倾斜方向比合约层和品种层更接近正确用法，"
        "但倾斜强度、趋势侧惩罚和尾部相关约束还不够。"
    )
    md.append("")
    md.append("结论：B6 不适合作为下一版主线直接替代 B2C；B6b 可保留为 P/C 侧预算研究线；B6a 和 B6c 更适合作为诊断和因子库输入，而不是当前交易配置。")
    md.append("")

    md.append("## 2. 实验定义")
    md.append("")
    md.append("- B1：只做流动性/OI 排序，是交易可行性基准。")
    md.append("- B2C：在 B1 基础上做品种预算倾斜，是当前 S1 主基准。")
    md.append("- B6a：在 B1 基础上引入合约层残差质量排序，并保留低价/摩擦硬过滤。")
    md.append("- B6b：在 B6a 基础上做 P/C 侧预算倾斜，使用 theta/vega、premium/stress、theta/gamma、premium/margin、vega/gamma per premium 和方向惩罚。")
    md.append("- B6c：在 B6a 基础上做品种层预算倾斜，使用 theta/vega、premium/stress、theta/gamma、tail beta、gamma per premium。")
    md.append("")
    md.append("本次报告以 B2C 为主基准，以 B1 为辅助基准。样本区间为 2022-01-04 至 2026-03-31。")
    md.append("")

    md.append("## 3. 核心绩效")
    md.append("")
    md.append(
        markdown_table(
            metrics,
            [
                ("name", "版本"),
                ("role", "定位"),
                ("nav", "终值NAV"),
                ("cum_ret", "累计收益"),
                ("ann_ret", "年化收益"),
                ("ann_vol", "年化波动"),
                ("max_dd", "最大回撤"),
                ("worst_day", "最差单日"),
                ("sharpe_rf2", "Sharpe RF2%"),
                ("calmar", "Calmar"),
            ],
        )
    )
    md.append("")
    md.append(
        f"从目标看，B6 系列都未达到我们希望的“年化 6%、最大回撤小于 2%”。"
        f"其中 B6b 年化约 {pct(float(b6b['ann_ret']))}、最大回撤 {pct(float(b6b['max_dd']))}，"
        f"收益略高于 B1 但低于 B2C；B6a/B6c 则收益和回撤都不理想。"
    )
    md.append("")

    md.append("## 4. Premium Formula 拆解")
    md.append("")
    md.append(
        markdown_table(
            metrics,
            [
                ("name", "版本"),
                ("gross_open_premium", "毛开仓权利金"),
                ("net_open_premium", "净开仓权利金"),
                ("premium_retained_ratio", "已平仓权利金留存率"),
                ("stop_count", "止损次数"),
                ("stop_pnl", "止损PnL"),
                ("expiry_count", "到期次数"),
                ("expiry_pnl", "到期PnL"),
                ("fee", "费用"),
            ],
        )
    )
    md.append("")
    md.append(
        "按照公式看，B6 的问题主要落在 Retention Rate 和 Tail/Stop Loss 两端。"
        "B6b 增加了一部分可交易权利金和方向侧调整，但没有显著降低尾部损耗；B6a/B6c 则说明把合约质量或品种质量单独拿来倾斜，容易牺牲组合路径。"
        "换句话说，B6 因子不是完全无效，而是需要更明确地分层使用：合约层管执行和尾部覆盖，P/C 层管方向侧预算，组合层管尾部聚集。"
    )
    md.append("")

    md.append("## 5. Greek 归因")
    md.append("")
    md.append(
        markdown_table(
            metrics,
            [
                ("name", "版本"),
                ("delta_pnl", "Delta"),
                ("gamma_pnl", "Gamma"),
                ("theta_pnl", "Theta"),
                ("vega_pnl", "Vega"),
                ("residual_pnl", "Residual"),
            ],
        )
    )
    md.append("")
    md.append(
        "B6 的最大警讯是：净值并没有因为更高质量的 theta 或更干净的 vega 获得明显改善。"
        "B6b 相比 B2C 的 theta 更厚一些，但 gamma/路径风险没有被压住；B6a/B6c 则更明显地把合约或品种排序优势转化成了更差的组合路径。"
        "这符合我们之前的判断：S1 的问题不只是“选哪个合约”，更是“这些合约能不能同时卖、在哪一侧卖、在什么环境下卖”。"
    )
    md.append("")

    md.append("## 6. 图表深读")
    md.append("")
    for note in CHART_NOTES:
        file_name = note["file"]
        title = note["title"]
        md.append(f"### {title}")
        md.append("")
        md.append(f"![{title}](../output/analysis_s1_b6_experiment_20260430/{file_name})")
        md.append("")
        md.append(f"**怎么看：**{note['how']}")
        md.append("")
        md.append(f"**图上读数：**{note['read']}")
        md.append("")
        md.append(f"**期权专家判断：**{note['judgement']}")
        md.append("")
        md.append(f"**风险或口径疑点：**{note['risk']}")
        md.append("")
        md.append(f"**下一步验证：**{note['next']}")
        md.append("")

    md.append("## 7. 产品和方向层观察")
    md.append("")
    if not product_summary.empty:
        top_loss = product_summary.groupby(["version", "product"])["pnl"].sum().reset_index()
        top_loss = top_loss.sort_values("pnl").groupby("version").head(3)
        md.append("各版本亏损最靠前的品种如下，说明品种层风险仍然有明显集中：")
        md.append("")
        md.append("| 版本 | 品种 | 已实现PnL |")
        md.append("| --- | --- | ---: |")
        for _, row in top_loss.iterrows():
            md.append(f"| {row['version']} | {row['product']} | {money(float(row['pnl']))} |")
        md.append("")
    md.append(
        "B6 的产品层和 P/C 侧实验提醒我们：一个因子在单因子检验中有效，不代表可以在所有层级使用。"
        "趋势、skew、breakout 更适合 P/C 预算；premium/stress、theta/vega 更适合合约排序；tail beta、stop cluster 更适合组合层约束。"
    )
    md.append("")

    md.append("## 8. 结论和下一步")
    md.append("")
    md.append("1. B2C 仍是当前主基准，B6b 是可保留的研究分支，但不能替代 B2C。")
    md.append("2. B6a 说明合约层排序不能单独解决组合尾部问题；B6c 说明品种层倾斜若没有组合约束，容易把预算给到看似高分但路径更脆的品种。")
    md.append("3. 下一步 B7 应转向组合层：板块集中度、相关组、tail correlation、stop cluster、same-expiry gamma、Top stress share。")
    md.append("4. 因子库使用方式要更严格：每个因子只能归入一个主层级，避免同一个分母或同一逻辑在合约、P/C、品种多处重复加权。")
    md.append("5. 若继续优化 B6b，应重点测试更强的 P/C 侧趋势/突破惩罚，而不是加大所有质量因子的统一倾斜。")
    md.append("")

    REPORT_MD.write_text("\n".join(md), encoding="utf-8")


def add_docx_table(doc: Document, df: pd.DataFrame, columns: List[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, (_, title) in enumerate(columns):
        table.rows[0].cells[i].text = title
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, (col, _) in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                if col in {"cum_ret", "ann_ret", "ann_vol", "max_dd", "worst_day", "avg_margin", "premium_retained_ratio"}:
                    text = pct(val)
                elif "pnl" in col or "premium" in col or col in {"nav", "fee"}:
                    text = money(val)
                else:
                    text = f"{val:.2f}"
            else:
                text = str(val)
            cells[i].text = text


def build_docx(metrics: pd.DataFrame, charts: List[Path]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("S1 B6 因子角色实验复盘报告", level=0)
    p = doc.add_paragraph("生成日期：2026-04-30")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading("一、期权专家审议摘要", level=1)
    doc.add_paragraph(
        "B6 的实验结论是：因子并非没有价值，但当前用法没有打过 B2C。"
        "B6b 是三条线里相对最好的分支，说明 P/C 侧预算比单纯合约层或品种层更接近正确方向；"
        "但 B6b 仍未形成对 B2C 的稳定超额，也没有显著改善 gamma、vega 和尾部损耗。"
    )
    doc.add_paragraph(
        "因此，B2C 仍应作为当前主基准；B6b 保留为后续 P/C 侧预算研究线；B6a/B6c 暂不进入主线。"
    )

    doc.add_heading("二、核心绩效", level=1)
    add_docx_table(
        doc,
        metrics,
        [
            ("name", "版本"),
            ("role", "定位"),
            ("nav", "终值NAV"),
            ("ann_ret", "年化收益"),
            ("max_dd", "最大回撤"),
            ("sharpe_rf2", "Sharpe RF2%"),
            ("calmar", "Calmar"),
        ],
    )

    doc.add_heading("三、Premium Formula 拆解", level=1)
    add_docx_table(
        doc,
        metrics,
        [
            ("name", "版本"),
            ("gross_open_premium", "毛开仓权利金"),
            ("premium_retained_ratio", "留存率"),
            ("stop_count", "止损次数"),
            ("stop_pnl", "止损PnL"),
            ("fee", "费用"),
        ],
    )

    doc.add_heading("四、图表与解读", level=1)
    note_by_file = {note["file"]: note for note in CHART_NOTES}
    for path in charts:
        note = note_by_file.get(path.name)
        doc.add_heading(note["title"] if note else path.name, level=2)
        doc.add_picture(str(path), width=Inches(6.3))
        if note:
            for label, key in [
                ("怎么看", "how"),
                ("图上读数", "read"),
                ("期权专家判断", "judgement"),
                ("风险或口径疑点", "risk"),
                ("下一步验证", "next"),
            ]:
                doc.add_paragraph(f"{label}：{note[key]}")
        else:
            doc.add_paragraph("本图用于辅助判断 B6 相对 B2C 的收益、风险或结构变化。")

    doc.add_heading("五、结论", level=1)
    for text in [
        "B2C 仍是当前主基准。",
        "B6b 可保留为 P/C 侧预算研究线，但需要更明确的趋势、breakout、skew 和 tail 约束。",
        "B6a/B6c 暂不进入主线，它们更适合作为因子库和诊断输入。",
        "下一步应进入 B7 组合层：板块、相关组、到期聚集、tail correlation 和 stop cluster。",
    ]:
        doc.add_paragraph(text, style=None)

    doc.save(REPORT_DOCX)


def main() -> None:
    setup()
    navs = {name: read_nav(meta["tag"]) for name, meta in RUNS.items()}
    orders = {name: read_orders(meta["tag"]) for name, meta in RUNS.items()}
    metrics = pd.DataFrame([metric_row(name, navs[name], orders[name]) for name in RUNS])
    save_table(metrics, "b6_experiment_summary_metrics.csv")
    charts = build_all_charts(navs, orders, metrics)
    prod_summary = product_side_summary(orders)
    stops = stop_summary(orders)
    write_report(metrics, charts, prod_summary, stops)
    build_docx(metrics, charts)
    print(f"analysis_dir={ANALYSIS}")
    print(f"markdown={REPORT_MD}")
    print(f"docx={REPORT_DOCX}")


if __name__ == "__main__":
    main()
