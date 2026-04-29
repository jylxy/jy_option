#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert an S1 factor-layer Markdown report into Word/Feishu DOCX.

The factor-layer reports are written in Markdown because that is easiest for
research iteration, but Feishu import needs real embedded images and Word tables.
This script supports the Markdown subset used by
``docs/s1_b2c_factor_layer_report_*.md`` and keeps the output deliberately
simple: headings, paragraphs, bullet lists, code blocks, tables and inline PNGs.
"""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(31, 78, 121)
BODY_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build Word and Feishu-import friendly DOCX from a factor-layer Markdown report."
    )
    parser.add_argument("--markdown", required=True, type=Path, help="Input Markdown report path.")
    parser.add_argument("--output", required=True, type=Path, help="Output Word DOCX path.")
    parser.add_argument(
        "--feishu-output",
        type=Path,
        default=None,
        help="Optional Feishu DOCX path. The file is copied from --output after generation.",
    )
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=Path.cwd(),
        help="Repository root used to resolve Markdown image references.",
    )
    return parser.parse_args()


def set_run_font(run, font_name: str = BODY_FONT, size: Optional[float] = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_default_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    styles = doc.styles
    for style_name in ("Normal", "Body Text", "List Bullet"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = BODY_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            style.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = BODY_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            style.font.color.rgb = ACCENT


def page_width_inches(doc: Document) -> float:
    section = doc.sections[-1]
    width = section.page_width - section.left_margin - section.right_margin
    return width / 914400


def add_inline_markdown_runs(paragraph, text: str, size: float = 10.5) -> None:
    """Add a small subset of inline Markdown: bold and inline code."""

    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(clean_inline(text[pos : match.start()]))
            set_run_font(run, size=size)
        token = match.group(0)
        content = token[2:-2] if token.startswith("**") else token[1:-1]
        run = paragraph.add_run(content)
        if token.startswith("**"):
            run.bold = True
            set_run_font(run, size=size)
        else:
            set_run_font(run, MONO_FONT, size=max(size - 0.5, 8.0))
            run.font.color.rgb = RGBColor(88, 88, 88)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(clean_inline(text[pos:]))
        set_run_font(run, size=size)


def clean_inline(text: str) -> str:
    return text.replace("\\|", "|").replace("&nbsp;", " ").strip()


def add_paragraph(doc: Document, text: str, style: Optional[str] = None, size: float = 10.5):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph)
    add_inline_markdown_runs(paragraph, text, size=size)
    return paragraph


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    if not lines:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_spacing(paragraph, before=2, after=8, line=1.05)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, MONO_FONT, size=9.0)
    run.font.color.rgb = RGBColor(70, 70, 70)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    p_pr.append(shd)


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").strip()
    return bool(stripped) and all(ch in "-:| " for ch in stripped)


def split_table_row(line: str) -> List[str]:
    return [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]


def parse_table(lines: Sequence[str]) -> List[List[str]]:
    rows = [split_table_row(line) for line in lines if not is_table_separator(line)]
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows]


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    # Factor reports contain wide numeric tables. Rendering them as real Word
    # tables often collapses columns into vertical text after Feishu import, so
    # we use readable row cards instead.
    add_table_as_cards(doc, rows)


def add_table_as_cards(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    headers = [clean_inline(cell) for cell in rows[0]]
    add_paragraph(doc, "（表格已转为逐行摘要，避免 Word/飞书导入时挤压成竖排。）", size=9.5)
    for row in rows[1:]:
        parts = []
        for col_idx, header in enumerate(headers):
            if col_idx >= len(row):
                continue
            value = clean_inline(str(row[col_idx]))
            if value:
                parts.append(f"{header}: {value}")
        if parts:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            set_paragraph_spacing(paragraph, after=4, line=1.08)
            add_inline_markdown_runs(paragraph, "；".join(parts), size=9.5)


def resolve_image(markdown_path: Path, repo_root: Path, image_ref: str) -> Optional[Path]:
    image_ref = image_ref.strip().strip("<>")
    candidate = Path(image_ref)
    candidates = []
    if candidate.is_absolute():
        candidates.append(candidate)
    else:
        candidates.append((markdown_path.parent / candidate).resolve())
        candidates.append((repo_root / image_ref).resolve())
    for path in candidates:
        if path.exists() and path.is_file():
            return path
    return None


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=4, after=2)
    run = paragraph.add_run()

    max_width = min(page_width_inches(doc), 6.7)
    try:
        with Image.open(image_path) as img:
            width_px, height_px = img.size
        aspect = height_px / max(width_px, 1)
        width = max_width
        if aspect > 0.82:
            width = min(max_width, 5.9)
    except Exception:
        width = max_width

    run.add_picture(str(image_path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, after=8)
        run = cap.add_run(caption)
        set_run_font(run, size=9.0)
        run.italic = True
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_heading(doc: Document, title: str, level: int) -> None:
    if level == 1:
        paragraph = doc.add_heading(clean_inline(title), level=0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=12)
        return
    paragraph = doc.add_heading(clean_inline(title), level=min(level - 1, 3))
    set_paragraph_spacing(paragraph, before=10 if level <= 2 else 6, after=4)


def add_cover_note(doc: Document, markdown_path: Path) -> None:
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(note, before=0, after=10)
    run = note.add_run("S1 卖权策略因子分层研究报告 | Word / 飞书导入版")
    set_run_font(run, size=10.0)
    run.font.color.rgb = RGBColor(110, 110, 110)

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(source, after=14)
    run = source.add_run(f"源文件：{markdown_path.name}")
    set_run_font(run, size=9.0)
    run.font.color.rgb = RGBColor(130, 130, 130)


def markdown_to_docx(markdown_path: Path, output_path: Path, repo_root: Path) -> None:
    doc = Document()
    set_default_styles(doc)

    lines = markdown_path.read_text(encoding="utf-8-sig").splitlines()
    idx = 0
    seen_title = False

    while idx < len(lines):
        raw_line = lines[idx].rstrip()
        line = raw_line.strip()

        if not line:
            idx += 1
            continue

        if line.startswith("```"):
            code_lines: List[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx].rstrip())
                idx += 1
            idx += 1 if idx < len(lines) else 0
            add_code_block(doc, code_lines)
            continue

        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].rstrip())
                idx += 1
            add_table(doc, parse_table(table_lines))
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if image_match:
            alt_text, image_ref = image_match.group(1).strip(), image_match.group(2).strip()
            image_path = resolve_image(markdown_path, repo_root, image_ref)
            if image_path is None:
                add_paragraph(doc, f"[图片缺失] {alt_text}: {image_ref}")
            else:
                add_image(doc, image_path, alt_text)
            idx += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            add_heading(doc, heading_match.group(2), level)
            if level == 1 and not seen_title:
                add_cover_note(doc, markdown_path)
                seen_title = True
            idx += 1
            continue

        if line.startswith("- "):
            add_paragraph(doc, line[2:], style="List Bullet")
            idx += 1
            continue

        add_paragraph(doc, line)
        idx += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    args = parse_args()
    markdown_path = args.markdown.resolve()
    repo_root = args.repo_root.resolve()
    output_path = args.output.resolve()
    markdown_to_docx(markdown_path, output_path, repo_root)
    if args.feishu_output:
        feishu_path = args.feishu_output.resolve()
        feishu_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(output_path, feishu_path)
    print(f"wrote {output_path}")
    if args.feishu_output:
        print(f"wrote {args.feishu_output.resolve()}")


if __name__ == "__main__":
    main()
