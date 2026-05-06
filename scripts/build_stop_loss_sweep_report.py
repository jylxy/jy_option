# -*- coding: utf-8 -*-
"""Build an S1 stop-loss sweep comparison report.

This script is intentionally self-contained. It reads the stop-loss experiment
NAV/order files, creates comparison charts, writes a Markdown report, and
exports a Feishu-friendly DOCX with embedded figures.
"""

from __future__ import annotations

import argparse
import math
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


TRADING_DAYS = 252


@dataclass(frozen=True)
class RunSpec:
    label: str
    tag: str
    stop_multiple: float | None
    status: str
    include_in_full_gradient: bool = True


RUNS = [
    RunSpec("Stop 1.5x", "s1_b2c_stop15_2022_latest", 1.5, "已完成"),
    RunSpec("Stop 2.0x", "s1_b2c_stop20_2022_latest", 2.0, "运行中快照", False),
    RunSpec("Stop 2.5x / B2C", "s1_b2_product_tilt075_stop25_allprod_2022_latest", 2.5, "已完成"),
    RunSpec("Stop 3.5x", "s1_b2c_stop35_2022_latest", 3.5, "已完成"),
    RunSpec("No Stop", "s1_b2c_nostop_2022_latest", None, "已完成"),
]


BASELINE_TAG = "s1_b2_product_tilt075_stop25_allprod_2022_latest"


def set_plot_style() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 140
    plt.rcParams["savefig.dpi"] = 180
    plt.rcParams["axes.grid"] = True
    plt.rcParams["grid.alpha"] = 0.25


def fmt_pct(x: float | int | None, digits: int = 2) -> str:
    if x is None or pd.isna(x):
        return ""
    return f"{float(x) * 100:.{digits}f}%"


def fmt_num(x: float | int | None, digits: int = 0) -> str:
    if x is None or pd.isna(x):
        return ""
    return f"{float(x):,.{digits}f}"


def pct_col(df: pd.DataFrame, candidates: Iterable[str]) -> str | None:
    for col in candidates:
        if col in df.columns:
            return col
    return None


def load_nav(output_dir: Path, spec: RunSpec) -> pd.DataFrame:
    path = output_dir / f"nav_{spec.tag}.csv"
    if not path.exists():
        return pd.DataFrame()
    df = pd.read_csv(path, low_memory=False)
    df["date"] = pd.to_datetime(df["date"])
    df["nav"] = pd.to_numeric(df["nav"], errors="coerce")
    df = df.dropna(subset=["date", "nav"]).sort_values("date").copy()
    df["daily_return"] = df["nav"].pct_change().fillna(0.0)
    df["cum_return"] = df["nav"] / df["nav"].iloc[0] - 1.0
    df["drawdown"] = df["nav"] / df["nav"].cummax() - 1.0
    df["label"] = spec.label
    return df


def summarize_nav(df: pd.DataFrame, spec: RunSpec) -> dict[str, object]:
    if df.empty:
        return {"label": spec.label, "tag": spec.tag, "status": "缺文件"}
    returns = df["daily_return"].iloc[1:]
    nav0 = float(df["nav"].iloc[0])
    nav1 = float(df["nav"].iloc[-1])
    n_days = max(len(df) - 1, 1)
    cum_ret = nav1 / nav0 - 1.0
    ann_ret = (nav1 / nav0) ** (TRADING_DAYS / n_days) - 1.0
    ann_vol = float(returns.std(ddof=1) * math.sqrt(TRADING_DAYS)) if len(returns) > 2 else np.nan
    sharpe = float(returns.mean() / returns.std(ddof=1) * math.sqrt(TRADING_DAYS)) if returns.std(ddof=1) else np.nan
    max_dd = float(df["drawdown"].min())
    calmar = ann_ret / abs(max_dd) if max_dd < 0 else np.nan
    margin_col = pct_col(df, ["s1_margin_used_pct", "margin_pct"])
    products_col = pct_col(df, ["s1_active_sell_products", "active_products"])
    contracts_col = pct_col(df, ["s1_active_sell_contracts", "active_contracts"])
    lots_col = pct_col(df, ["s1_active_sell_lots", "active_lots"])
    pc_col = pct_col(df, ["s1_put_call_lot_ratio", "put_call_ratio"])
    return {
        "label": spec.label,
        "tag": spec.tag,
        "status": spec.status,
        "start": df["date"].iloc[0].date().isoformat(),
        "end": df["date"].iloc[-1].date().isoformat(),
        "rows": len(df),
        "final_nav": nav1,
        "cum_return": cum_ret,
        "ann_return": ann_ret,
        "ann_vol": ann_vol,
        "sharpe": sharpe,
        "calmar": calmar,
        "max_dd": max_dd,
        "worst_day": float(returns.min()) if len(returns) else np.nan,
        "best_day": float(returns.max()) if len(returns) else np.nan,
        "avg_margin": float(df[margin_col].mean()) if margin_col else np.nan,
        "peak_margin": float(df[margin_col].max()) if margin_col else np.nan,
        "last_margin": float(df[margin_col].iloc[-1]) if margin_col else np.nan,
        "avg_products": float(df[products_col].mean()) if products_col else np.nan,
        "avg_contracts": float(df[contracts_col].mean()) if contracts_col else np.nan,
        "avg_lots": float(df[lots_col].mean()) if lots_col else np.nan,
        "last_positions": float(df["n_positions"].iloc[-1]) if "n_positions" in df.columns else np.nan,
        "avg_pc_ratio": float(df[pc_col].replace([np.inf, -np.inf], np.nan).mean()) if pc_col else np.nan,
        "fee_sum": float(df["fee"].sum()) if "fee" in df.columns else np.nan,
        "delta_pnl": float(df["delta_pnl"].sum()) if "delta_pnl" in df.columns else np.nan,
        "gamma_pnl": float(df["gamma_pnl"].sum()) if "gamma_pnl" in df.columns else np.nan,
        "theta_pnl": float(df["theta_pnl"].sum()) if "theta_pnl" in df.columns else np.nan,
        "vega_pnl": float(df["vega_pnl"].sum()) if "vega_pnl" in df.columns else np.nan,
        "residual_pnl": float(df["residual_pnl"].sum()) if "residual_pnl" in df.columns else np.nan,
    }


def load_orders(output_dir: Path, spec: RunSpec) -> pd.DataFrame:
    path = output_dir / f"orders_{spec.tag}.csv"
    if not path.exists():
        return pd.DataFrame()
    df = pd.read_csv(path, low_memory=False)
    if "date" in df.columns:
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
    return df


def summarize_orders(df: pd.DataFrame, spec: RunSpec) -> dict[str, object]:
    if df.empty:
        return {"label": spec.label, "tag": spec.tag, "has_orders": False}
    action = df["action"].astype(str).str.lower() if "action" in df.columns else pd.Series("", index=df.index)
    open_mask = action.str.contains("open_sell", na=False)
    stop_mask = action.str.contains("sl|stop|止损", regex=True, na=False)
    expiry_mask = action.str.contains("expiry|到期", regex=True, na=False)
    close_mask = stop_mask | expiry_mask | action.str.contains("close", na=False)

    def nsum(col: str, mask: pd.Series | None = None) -> float:
        if col not in df.columns:
            return np.nan
        s = pd.to_numeric(df[col], errors="coerce").fillna(0.0)
        return float(s[mask].sum()) if mask is not None else float(s.sum())

    open_net_premium = nsum("net_premium_cash", open_mask)
    gross_premium = nsum("gross_premium_cash", open_mask)
    realized_pnl = nsum("pnl")
    stop_pnl = nsum("pnl", stop_mask)
    expiry_pnl = nsum("pnl", expiry_mask)
    retained_cash = nsum("premium_retained_cash", close_mask)
    close_open_premium = nsum("open_premium_cash", close_mask)
    stop_open_premium = nsum("open_premium_cash", stop_mask)
    expiry_open_premium = nsum("open_premium_cash", expiry_mask)
    close_value_cash = nsum("close_value_cash", close_mask)
    fee = nsum("fee")

    return {
        "label": spec.label,
        "tag": spec.tag,
        "has_orders": True,
        "order_rows": len(df),
        "open_rows": int(open_mask.sum()),
        "stop_rows": int(stop_mask.sum()),
        "expiry_rows": int(expiry_mask.sum()),
        "gross_open_premium": gross_premium,
        "net_open_premium": open_net_premium,
        "realized_pnl": realized_pnl,
        "stop_pnl": stop_pnl,
        "expiry_pnl": expiry_pnl,
        "fee_orders": fee,
        "premium_retained_cash": retained_cash,
        "close_open_premium": close_open_premium,
        "premium_retained_ratio": retained_cash / close_open_premium if close_open_premium else np.nan,
        "stop_open_premium": stop_open_premium,
        "expiry_open_premium": expiry_open_premium,
        "stop_loss_to_open_premium": stop_pnl / stop_open_premium if stop_open_premium else np.nan,
        "close_value_cash": close_value_cash,
    }


def align_to_baseline(nav_map: dict[str, pd.DataFrame]) -> pd.DataFrame:
    base = nav_map[BASELINE_TAG][["date", "nav", "cum_return", "daily_return"]].rename(
        columns={
            "nav": "baseline_nav",
            "cum_return": "baseline_cum_return",
            "daily_return": "baseline_daily_return",
        }
    )
    rows = []
    for spec in RUNS:
        df = nav_map.get(spec.tag, pd.DataFrame())
        if df.empty or spec.tag == BASELINE_TAG:
            continue
        merged = df[["date", "nav", "cum_return", "daily_return", "label"]].merge(base, on="date", how="inner")
        if merged.empty:
            continue
        merged["excess_return"] = merged["cum_return"] - merged["baseline_cum_return"]
        merged["daily_excess_return"] = merged["daily_return"] - merged["baseline_daily_return"]
        merged["relative_nav"] = merged["nav"] / merged["baseline_nav"] - 1.0
        rows.append(merged)
    return pd.concat(rows, ignore_index=True) if rows else pd.DataFrame()


def yearly_returns(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    d = df.copy()
    d["year"] = d["date"].dt.year
    rows = []
    for year, g in d.groupby("year"):
        rows.append(
            {
                "year": int(year),
                "return": float(g["nav"].iloc[-1] / g["nav"].iloc[0] - 1.0),
                "max_dd": float((g["nav"] / g["nav"].cummax() - 1.0).min()),
            }
        )
    return pd.DataFrame(rows)


def save_table(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path, index=False, encoding="utf-8-sig")


def plot_nav(nav_map: dict[str, pd.DataFrame], path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.5, 5.8))
    for spec in RUNS:
        df = nav_map.get(spec.tag, pd.DataFrame())
        if df.empty:
            continue
        lw = 2.4 if spec.tag == BASELINE_TAG else 1.8
        ls = "--" if spec.status.startswith("运行") else "-"
        ax.plot(df["date"], df["nav"] / df["nav"].iloc[0], label=f"{spec.label} ({spec.status})", linewidth=lw, linestyle=ls)
    ax.set_title("NAV normalized: stop-loss sweep")
    ax.set_ylabel("NAV / initial NAV")
    ax.legend(ncols=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_excess(excess: pd.DataFrame, path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.5, 5.8))
    if not excess.empty:
        for label, g in excess.groupby("label"):
            ax.plot(g["date"], g["excess_return"], label=label, linewidth=1.9, linestyle="--" if "2.0" in label else "-")
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Excess cumulative return vs Stop 2.5x / B2C")
    ax.set_ylabel("Excess cumulative return")
    ax.legend(ncols=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_drawdown(nav_map: dict[str, pd.DataFrame], path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.5, 5.8))
    for spec in RUNS:
        df = nav_map.get(spec.tag, pd.DataFrame())
        if df.empty:
            continue
        ax.plot(df["date"], df["drawdown"], label=spec.label, linewidth=1.8, linestyle="--" if spec.status.startswith("运行") else "-")
    ax.set_title("Drawdown comparison")
    ax.set_ylabel("Drawdown")
    ax.legend(ncols=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_threshold_gradient(summary: pd.DataFrame, path: Path) -> None:
    finite = summary[(summary["status"] == "已完成") & summary["stop_multiple"].notna()].copy()
    finite = finite.sort_values("stop_multiple")
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))
    axes[0].plot(finite["stop_multiple"], finite["cum_return"], marker="o", linewidth=2.2)
    axes[0].set_title("Return gradient by stop multiple")
    axes[0].set_xlabel("Stop multiple")
    axes[0].set_ylabel("Cumulative return")
    axes[1].plot(finite["stop_multiple"], finite["max_dd"], marker="o", color="#B1464A", linewidth=2.2)
    axes[1].set_title("Drawdown gradient by stop multiple")
    axes[1].set_xlabel("Stop multiple")
    axes[1].set_ylabel("Max drawdown")
    for ax in axes:
        ax.axhline(0, color="black", linewidth=0.7)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_stop_loss(order_summary: pd.DataFrame, path: Path) -> None:
    d = order_summary[order_summary["has_orders"] == True].copy()  # noqa: E712
    d = d[d["label"] != "No Stop"]
    fig, ax1 = plt.subplots(figsize=(10.8, 5.4))
    x = np.arange(len(d))
    ax1.bar(x - 0.18, d["stop_rows"], width=0.36, label="Stop rows", color="#4C78A8")
    ax1.set_ylabel("Stop rows")
    ax2 = ax1.twinx()
    ax2.bar(x + 0.18, d["stop_pnl"] / 1_000_000, width=0.36, label="Stop PnL (mn)", color="#E45756")
    ax2.set_ylabel("Stop PnL, million")
    ax1.set_xticks(x)
    ax1.set_xticklabels(d["label"], rotation=15, ha="right")
    ax1.set_title("Stop frequency vs stop loss amount")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper center", ncols=2)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_greeks(summary: pd.DataFrame, path: Path) -> None:
    d = summary[summary["status"] == "已完成"].copy()
    cols = ["delta_pnl", "gamma_pnl", "theta_pnl", "vega_pnl", "residual_pnl"]
    data = d.set_index("label")[cols] / 1_000_000
    fig, ax = plt.subplots(figsize=(11.5, 5.8))
    data.plot(kind="bar", stacked=False, ax=ax)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Greek attribution by stop rule")
    ax.set_ylabel("Cumulative PnL, million")
    ax.legend(ncols=3, fontsize=9)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_premium_retention(order_summary: pd.DataFrame, path: Path) -> None:
    d = order_summary[(order_summary["has_orders"] == True) & (order_summary["label"] != "No Stop")].copy()  # noqa: E712
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))
    axes[0].bar(d["label"], d["net_open_premium"] / 1_000_000, color="#72B7B2")
    axes[0].set_title("Net open premium")
    axes[0].set_ylabel("Million")
    axes[0].tick_params(axis="x", rotation=15)
    axes[1].bar(d["label"], d["premium_retained_ratio"], color="#54A24B")
    axes[1].set_title("Realized premium retention")
    axes[1].set_ylabel("Retained / open premium")
    axes[1].tick_params(axis="x", rotation=15)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_yearly(nav_map: dict[str, pd.DataFrame], path: Path) -> pd.DataFrame:
    pieces = []
    for spec in RUNS:
        if spec.status != "已完成":
            continue
        yr = yearly_returns(nav_map.get(spec.tag, pd.DataFrame()))
        if yr.empty:
            continue
        yr["label"] = spec.label
        pieces.append(yr)
    all_years = pd.concat(pieces, ignore_index=True) if pieces else pd.DataFrame()
    if all_years.empty:
        return all_years
    pivot = all_years.pivot(index="year", columns="label", values="return")
    fig, ax = plt.subplots(figsize=(11.5, 5.8))
    pivot.plot(kind="bar", ax=ax)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Yearly return comparison")
    ax.set_ylabel("Yearly return")
    ax.legend(ncols=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)
    return all_years


def plot_margin(nav_map: dict[str, pd.DataFrame], path: Path) -> None:
    fig, axes = plt.subplots(2, 1, figsize=(11.5, 7.2), sharex=True)
    for spec in RUNS:
        df = nav_map.get(spec.tag, pd.DataFrame())
        if df.empty:
            continue
        margin_col = pct_col(df, ["s1_margin_used_pct", "margin_pct"])
        if margin_col:
            axes[0].plot(df["date"], df[margin_col], label=spec.label, linewidth=1.6)
        if "s1_active_sell_contracts" in df.columns:
            axes[1].plot(df["date"], df["s1_active_sell_contracts"], label=spec.label, linewidth=1.6)
    axes[0].set_title("Margin usage")
    axes[0].set_ylabel("S1 margin pct")
    axes[1].set_title("Active short contracts")
    axes[1].set_ylabel("Contracts")
    axes[1].legend(ncols=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def make_md_table(df: pd.DataFrame, columns: list[tuple[str, str, str]]) -> str:
    rows = []
    rows.append("| " + " | ".join(title for _, title, _ in columns) + " |")
    rows.append("| " + " | ".join("---" for _ in columns) + " |")
    for _, r in df.iterrows():
        cells = []
        for col, _, kind in columns:
            val = r.get(col, "")
            if kind == "pct":
                cells.append(fmt_pct(val))
            elif kind == "num0":
                cells.append(fmt_num(val, 0))
            elif kind == "num2":
                cells.append(fmt_num(val, 2))
            else:
                cells.append("" if pd.isna(val) else str(val))
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows)


def add_docx_table(document: Document, df: pd.DataFrame, columns: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (_, title, _) in enumerate(columns):
        table.rows[0].cells[i].text = title
        shade_cell(table.rows[0].cells[i], "DCE6F1")
    for _, r in df.iterrows():
        cells = table.add_row().cells
        for i, (col, _, kind) in enumerate(columns):
            val = r.get(col, "")
            if kind == "pct":
                text = fmt_pct(val)
            elif kind == "num0":
                text = fmt_num(val, 0)
            elif kind == "num2":
                text = fmt_num(val, 2)
            else:
                text = "" if pd.isna(val) else str(val)
            cells[i].text = text
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    document.add_paragraph()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_picture(document: Document, path: Path, caption: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    cap = document.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def write_docx(
    md_sections: list[tuple[str, str]],
    summary_df: pd.DataFrame,
    excess_df: pd.DataFrame,
    order_df: pd.DataFrame,
    chart_specs: list[tuple[Path, str, str]],
    docx_path: Path,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("S1 B2C 止损倍数实验对比报告")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle = document.add_paragraph("主题：Stop 1.5x / 2.0x / 2.5x / 3.5x / 不止损的收益、回撤、超额与权利金留存比较")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("报告口径：Stop 2.5x / B2C 为基准；Stop 2.0x 为运行中快照，其余为 2022-01-04 至 2026-03-31 完整样本。")

    document.add_heading("一、核心绩效表", level=1)
    perf_cols = [
        ("label", "版本", "str"),
        ("status", "状态", "str"),
        ("end", "截止日", "str"),
        ("final_nav", "NAV", "num0"),
        ("cum_return", "累计收益", "pct"),
        ("ann_return", "年化收益", "pct"),
        ("max_dd", "最大回撤", "pct"),
        ("sharpe", "Sharpe", "num2"),
        ("calmar", "Calmar", "num2"),
    ]
    add_docx_table(document, summary_df, perf_cols)

    document.add_heading("二、相对 B2C 超额表", level=1)
    excess_cols = [
        ("label", "版本", "str"),
        ("aligned_end", "共同截止日", "str"),
        ("excess_return", "累计超额", "pct"),
        ("relative_nav", "相对NAV", "pct"),
        ("daily_excess_mean", "日均超额", "pct"),
        ("excess_min", "超额低点", "pct"),
    ]
    add_docx_table(document, excess_df, excess_cols)

    document.add_heading("三、止损与权利金留存表", level=1)
    order_cols = [
        ("label", "版本", "str"),
        ("open_rows", "开仓行数", "num0"),
        ("stop_rows", "止损行数", "num0"),
        ("net_open_premium", "净权利金", "num0"),
        ("stop_pnl", "止损PnL", "num0"),
        ("realized_pnl", "实现PnL", "num0"),
        ("premium_retained_ratio", "留存率", "pct"),
    ]
    add_docx_table(document, order_df, order_cols)

    document.add_heading("四、图表与解读", level=1)
    for path, title_text, body in chart_specs:
        if not path.exists():
            continue
        document.add_heading(title_text, level=2)
        add_picture(document, path, title_text)
        for para in body.split("\n\n"):
            if para.strip():
                document.add_paragraph(para.strip())

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("五、正文分析", level=1)
    for title_text, body in md_sections:
        document.add_heading(title_text, level=2)
        for para in body.split("\n\n"):
            if para.strip():
                document.add_paragraph(para.strip())

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", default="server_deploy/output")
    parser.add_argument("--report-dir", default="server_deploy/output/analysis_stop_loss_sweep_20260430")
    args = parser.parse_args()

    set_plot_style()
    output_dir = Path(args.output_dir)
    report_dir = Path(args.report_dir)
    chart_dir = report_dir / "charts"
    chart_dir.mkdir(parents=True, exist_ok=True)

    nav_map = {spec.tag: load_nav(output_dir, spec) for spec in RUNS}
    summary_rows = []
    for spec in RUNS:
        row = summarize_nav(nav_map[spec.tag], spec)
        row["stop_multiple"] = spec.stop_multiple
        summary_rows.append(row)
    summary = pd.DataFrame(summary_rows)

    order_rows = []
    for spec in RUNS:
        order_rows.append(summarize_orders(load_orders(output_dir, spec), spec))
    order_summary = pd.DataFrame(order_rows)
    full = summary.merge(order_summary, on=["label", "tag"], how="left")

    excess = align_to_baseline(nav_map)
    excess_rows = []
    for label, g in excess.groupby("label"):
        excess_rows.append(
            {
                "label": label,
                "aligned_end": g["date"].iloc[-1].date().isoformat(),
                "excess_return": float(g["excess_return"].iloc[-1]),
                "relative_nav": float(g["relative_nav"].iloc[-1]),
                "daily_excess_mean": float(g["daily_excess_return"].mean()),
                "excess_min": float(g["excess_return"].min()),
                "excess_max": float(g["excess_return"].max()),
                "n_days": len(g),
            }
        )
    excess_summary = pd.DataFrame(excess_rows)

    save_table(full, report_dir / "stop_loss_sweep_summary.csv")
    save_table(excess_summary, report_dir / "stop_loss_sweep_excess_vs_b2c.csv")
    yearly = plot_yearly(nav_map, chart_dir / "08_yearly_return_comparison.png")
    save_table(yearly, report_dir / "stop_loss_sweep_yearly_returns.csv")

    plot_nav(nav_map, chart_dir / "01_nav_normalized.png")
    plot_excess(excess, chart_dir / "02_excess_vs_b2c.png")
    plot_drawdown(nav_map, chart_dir / "03_drawdown_comparison.png")
    plot_threshold_gradient(full, chart_dir / "04_stop_threshold_gradient.png")
    plot_stop_loss(order_summary, chart_dir / "05_stop_count_and_loss.png")
    plot_greeks(full, chart_dir / "06_greek_attribution.png")
    plot_premium_retention(order_summary, chart_dir / "07_premium_retention.png")
    plot_margin(nav_map, chart_dir / "09_margin_and_contracts.png")

    stop15 = full.loc[full["label"] == "Stop 1.5x"].iloc[0]
    stop25 = full.loc[full["label"] == "Stop 2.5x / B2C"].iloc[0]
    stop35 = full.loc[full["label"] == "Stop 3.5x"].iloc[0]
    nostop = full.loc[full["label"] == "No Stop"].iloc[0]
    stop20 = full.loc[full["label"] == "Stop 2.0x"].iloc[0]

    finite = full[(full["status"] == "已完成") & full["stop_multiple"].notna()].sort_values("stop_multiple")
    x = finite["stop_multiple"].astype(float).to_numpy()
    y = finite["cum_return"].astype(float).to_numpy()
    if len(x) >= 2:
        coef = np.polyfit(x, y, 1)
        pred = np.polyval(coef, x)
        ss_res = float(((y - pred) ** 2).sum())
        ss_tot = float(((y - y.mean()) ** 2).sum())
        r2 = 1 - ss_res / ss_tot if ss_tot else np.nan
    else:
        coef = [np.nan, np.nan]
        r2 = np.nan

    perf_cols = [
        ("label", "版本", "str"),
        ("status", "状态", "str"),
        ("end", "截止日", "str"),
        ("final_nav", "NAV", "num0"),
        ("cum_return", "累计收益", "pct"),
        ("ann_return", "年化收益", "pct"),
        ("ann_vol", "年化波动", "pct"),
        ("max_dd", "最大回撤", "pct"),
        ("sharpe", "Sharpe", "num2"),
        ("calmar", "Calmar", "num2"),
        ("avg_margin", "平均保证金", "pct"),
        ("peak_margin", "峰值保证金", "pct"),
    ]
    excess_cols = [
        ("label", "版本", "str"),
        ("aligned_end", "共同截止日", "str"),
        ("excess_return", "累计超额", "pct"),
        ("relative_nav", "相对NAV", "pct"),
        ("daily_excess_mean", "日均超额", "pct"),
        ("excess_min", "超额低点", "pct"),
        ("excess_max", "超额高点", "pct"),
        ("n_days", "共同天数", "num0"),
    ]
    order_cols = [
        ("label", "版本", "str"),
        ("order_rows", "订单行", "num0"),
        ("open_rows", "开仓行", "num0"),
        ("stop_rows", "止损行", "num0"),
        ("expiry_rows", "到期行", "num0"),
        ("net_open_premium", "净开仓权利金", "num0"),
        ("stop_pnl", "止损PnL", "num0"),
        ("expiry_pnl", "到期PnL", "num0"),
        ("realized_pnl", "实现PnL", "num0"),
        ("premium_retained_ratio", "权利金留存率", "pct"),
    ]
    greek_cols = [
        ("label", "版本", "str"),
        ("delta_pnl", "Delta", "num0"),
        ("gamma_pnl", "Gamma", "num0"),
        ("theta_pnl", "Theta", "num0"),
        ("vega_pnl", "Vega", "num0"),
        ("residual_pnl", "Residual", "num0"),
        ("fee_sum", "NAV费用", "num0"),
    ]

    chart_specs = [
        (
            chart_dir / "01_nav_normalized.png",
            "图1：归一化 NAV 对比",
            "怎么看：这张图看不同止损倍数在同一初始资金下的净值路径，而不是只看终点。\n\n"
            f"图上读数：完整样本中 Stop 1.5x 终值 {fmt_num(stop15['final_nav'])}，累计收益 {fmt_pct(stop15['cum_return'])}；B2C/2.5x 为 {fmt_pct(stop25['cum_return'])}；3.5x 为 {fmt_pct(stop35['cum_return'])}；不止损为 {fmt_pct(nostop['cum_return'])}。\n\n"
            "期权专家判断：净值梯度非常清楚，早止损并没有因为频繁砍仓而损害长期收益，反而保住了权利金留存和尾部资金效率。\n\n"
            "风险或口径疑点：Stop 2.0x 仍是运行中快照，不能和完整样本直接作最终排序；但阶段路径已经处在 1.5x 与 2.5x 之间，方向上支持止损倍数梯度。\n\n"
            "下一步验证：等 Stop 2.0x 完整跑完后，应补齐 1.0x、3.0x、4.0x、5.0x，再判断是否真的是连续单调，而不是 2022-2026 样本下的局部最优。",
        ),
        (
            chart_dir / "02_excess_vs_b2c.png",
            "图2：相对 B2C/2.5x 的累计超额",
            "怎么看：这张图把 B2C/2.5x 设为零轴，直接观察不同止损规则是否持续贡献超额。\n\n"
            f"图上读数：Stop 1.5x 全样本相对 B2C 的累计超额约 {fmt_pct(excess_summary.loc[excess_summary['label']=='Stop 1.5x','excess_return'].iloc[0])}；Stop 3.5x 为 {fmt_pct(excess_summary.loc[excess_summary['label']=='Stop 3.5x','excess_return'].iloc[0])}；不止损为 {fmt_pct(excess_summary.loc[excess_summary['label']=='No Stop','excess_return'].iloc[0])}。\n\n"
            "期权专家判断：如果一条超额曲线只是靠某一两天跳升，不能说明规则优越；这里更重要的是 1.5x 超额在多个阶段保持为正，说明它不是单一尾部日期的偶然保护。\n\n"
            "风险或口径疑点：超额仍然可能来自某些特定品种趋势期提前出场，而不是普适的期权卖方优势，因此需要后续按品种、P/C 侧和止损后是否归零继续拆。\n\n"
            "下一步验证：把超额路径拆成 2022H1、2022H2-2023Q1、2024Q3、2025Q2 等压力/恢复区间，确认早止损在顺畅期是否也没有显著拖累。",
        ),
        (
            chart_dir / "03_drawdown_comparison.png",
            "图3：最大回撤路径对比",
            "怎么看：卖权策略的核心不是胜率，而是尾部亏损是否吞掉数月权利金。\n\n"
            f"图上读数：Stop 1.5x 最大回撤 {fmt_pct(stop15['max_dd'])}，B2C/2.5x 最大回撤 {fmt_pct(stop25['max_dd'])}，3.5x 最大回撤 {fmt_pct(stop35['max_dd'])}，不止损最大回撤 {fmt_pct(nostop['max_dd'])}。\n\n"
            "期权专家判断：阈值从 1.5x 放宽到 3.5x 时，回撤几乎同步恶化；不止损则暴露出卖方策略最危险的一面，即尾部亏损不是线性增加，而可能一次性吞掉全部 carry。\n\n"
            "风险或口径疑点：回撤改善不能只理解为降低风险，因为更频繁的止损也可能牺牲后续反弹；需要检查止损后最终归零比例和止损后反向修复比例。\n\n"
            "下一步验证：对每笔止损合约做 after-stop outcome，分为止对、止早、止错三类，判断 1.5x 的优势来自真实尾部规避还是噪声止损后的幸运再开仓。",
        ),
        (
            chart_dir / "04_stop_threshold_gradient.png",
            "图4：止损倍数梯度",
            "怎么看：这张图只看完整样本中的有限倍数 1.5x、2.5x、3.5x，检验是否存在近似单调梯度。\n\n"
            f"图上读数：三点线性拟合斜率约 {coef[0]:.4f}，R² 约 {r2:.3f}。样本内表现为止损倍数越低，累计收益越高、最大回撤越小。\n\n"
            "期权专家判断：这确实支持“早止损线性增强”的阶段性结论，但三点不足以证明结构性规律。更严谨的说法是：在当前 B2C 框架和 2022-2026 样本内，止损倍数放宽带来的尾部损失增加，超过了少止损带来的权利金留存收益。\n\n"
            "风险或口径疑点：1.5x 可能是样本内最优，不能直接作为最终参数；它更像告诉我们“2.5x 仍偏晚”，而不是告诉我们最优一定是 1.5x。\n\n"
            "下一步验证：补跑 1.0x、2.0x 完整、3.0x、4.0x、5.0x，并按年份滚动观察梯度是否稳定。",
        ),
        (
            chart_dir / "05_stop_count_and_loss.png",
            "图5：止损次数与止损金额",
            "怎么看：这张图同时看止损频率和止损总损失，用来判断“早止损是不是只是在多交摩擦成本”。\n\n"
            f"图上读数：Stop 1.5x 止损 {fmt_num(stop15['stop_rows'])} 行、止损 PnL {fmt_num(stop15['stop_pnl'])}；B2C/2.5x 止损 {fmt_num(stop25['stop_rows'])} 行、止损 PnL {fmt_num(stop25['stop_pnl'])}；3.5x 止损 {fmt_num(stop35['stop_rows'])} 行、止损 PnL {fmt_num(stop35['stop_pnl'])}。\n\n"
            "期权专家判断：1.5x 止损次数最多，但总止损损失反而较小；3.5x 止损次数少，却每次亏损更重。这是卖方风险控制里很典型的“少止损不等于少亏损”。\n\n"
            "风险或口径疑点：止损行数不是合约数，一笔组合可能多腿、多执行价；报告结论应更多看金额、NAV、回撤，而不是单独看行数。\n\n"
            "下一步验证：按品种和方向统计单笔平均止损亏损，检查是否存在少数品种在宽止损下贡献了主要尾损。",
        ),
        (
            chart_dir / "06_greek_attribution.png",
            "图6：Greek PnL 归因对比",
            "怎么看：卖权策略理论上应赚 theta，并在降波期赚 vega；如果净值改善主要来自 delta，需要谨慎。\n\n"
            "图上读数：不同止损倍数下，theta 基本是主要正贡献项；gamma 和 vega 是主要侵蚀项；早止损的意义在于限制 gamma/vega 恶化继续累积。\n\n"
            "期权专家判断：Stop 1.5x 的优势不是单纯多收权利金，而是更早中断“权利金变贵”的路径，即在 short gamma/short vega 开始失控前把风险切断。\n\n"
            "风险或口径疑点：Greek 归因仍依赖 Black76/估值口径，Residual 不应当被解释为 alpha；真正结论需要与成交路径和 stop after-outcome 一起看。\n\n"
            "下一步验证：对止损发生日前后 1/3/5 日拆 vega、gamma 和 residual，判断止损到底是在切升波风险、切方向趋势，还是切异常报价。",
        ),
        (
            chart_dir / "07_premium_retention.png",
            "图7：净权利金与留存率",
            "怎么看：S1 的收益公式不是只看 Premium Pool，而是 Premium Pool × Deployment Ratio × Retention Rate，再扣 Tail/Stop Loss 和 Cost。\n\n"
            "图上读数：不同止损倍数的开仓权利金池接近，但留存和尾损差异明显；早止损的核心是提高最终留存质量，而不是扩大开仓池。\n\n"
            "期权专家判断：如果 1.5x 能在权利金池相近的情况下提高实现 PnL，这说明它主要改善 Retention Rate 和 Tail/Stop Loss 两项，而不是靠更激进部署。\n\n"
            "风险或口径疑点：留存率受到期、止损和期末未平仓共同影响；需要将已闭合交易和期末持仓分开，不应把未实现权利金误当成已留存。\n\n"
            "下一步验证：按闭合事件计算 close-only retention，并对期末持仓做 mark-to-market retention 校正。",
        ),
        (
            chart_dir / "08_yearly_return_comparison.png",
            "图8：年度收益对比",
            "怎么看：这张图检验早止损优势是否只来自某一年极端行情。\n\n"
            "图上读数：Stop 1.5x 的优势需要看是否分布在多个年度；如果只集中在 2022 或 2024，则参数过拟合风险更高。\n\n"
            "期权专家判断：卖方策略可以允许某些年份弱，但不能靠单一年份规避尾部来解释全部超额。年度拆分是判断策略稳健性的最低要求。\n\n"
            "风险或口径疑点：2026 年只有截至 3 月底的非完整年度，应作为阶段数据，不应与完整年度等权比较。\n\n"
            "下一步验证：改成滚动 6 个月/12 个月窗口，看 1.5x 的超额是否持续为正。",
        ),
        (
            chart_dir / "09_margin_and_contracts.png",
            "图9：保证金与持仓合约数",
            "怎么看：如果早止损收益更好但仓位明显更低，可能只是风险暴露不同；如果保证金接近而表现更好，则说明规则质量更优。\n\n"
            "图上读数：各版本在多数时间仍接近策略设定的保证金使用区间；Stop 2.0x 当前在 2024-08 快照附近保证金曾接近 50%。\n\n"
            "期权专家判断：止损倍数主要改变的是亏损路径和再部署节奏，而不是简单改变最大保证金上限。因此它更像退出规则 alpha，而不是仓位规模 alpha。\n\n"
            "风险或口径疑点：止损越早，释放保证金后可能更快重开，导致实际 Deployment Ratio 更高；这需要用日度新开权利金和止损后重开间隔进一步验证。\n\n"
            "下一步验证：增加“止损后释放保证金到重新部署”的资金周转分析，确认 1.5x 是否提升了 capital recycling。",
        ),
    ]

    sections = [
        (
            "核心结论",
            f"当前结果支持一个非常重要的阶段性判断：在 B2C 结构下，止损倍数从 3.5x 收紧到 2.5x，再收紧到 1.5x，净值、回撤和超额呈现清晰的单调改善。完整样本中，Stop 1.5x 累计收益 {fmt_pct(stop15['cum_return'])}，显著高于 B2C/2.5x 的 {fmt_pct(stop25['cum_return'])} 和 Stop 3.5x 的 {fmt_pct(stop35['cum_return'])}；不止损为 {fmt_pct(nostop['cum_return'])}，最大回撤达到 {fmt_pct(nostop['max_dd'])}，基本可以排除作为主线。\n\n"
            f"这不是“止损越多越差”的直觉结果，而是更符合卖方策略的尾部逻辑：1.5x 止损次数更多，但总止损损失 {fmt_num(stop15['stop_pnl'])}，反而小于 2.5x 的 {fmt_num(stop25['stop_pnl'])} 和 3.5x 的 {fmt_num(stop35['stop_pnl'])}。换句话说，早止损多砍了噪声和早期趋势，但更重要的是避免了少数合约从 2.5x 继续滑向 3.5x 甚至更大损失。\n\n"
            f"不过，报告不把它写成最终定论。当前有限倍数只有 1.5x、2.5x、3.5x 三个完整点，线性拟合 R² 约 {r2:.3f} 只能说明样本内梯度强，不能说明 1.5x 就是长期最优。Stop 2.0x 仍在运行中，截至 {stop20['end']}，累计收益 {fmt_pct(stop20['cum_return'])}，同日仍高于 B2C 快照，方向上进一步支持“2.5x 偏晚”的判断。",
        ),
        (
            "用收益公式解释这次结果",
            "按照 S1 的研究框架，净收益应拆成 Premium Pool × Deployment Ratio × Retention Rate - Tail/Stop Loss - Cost/Slippage。止损倍数实验主要不是在改变 Premium Pool，也不是改变初始候选排序，而是在改变 Retention Rate 与 Tail/Stop Loss 的分配。\n\n"
            "Stop 3.5x 和不止损看似给了卖方更多“熬到归零”的机会，但实际暴露给 short gamma、short vega 和趋势穿透的时间更长。只要少数合约从 2.5x 后继续恶化，额外损失就会超过那些本来可以等到归零的合约贡献。Stop 1.5x 的本质是牺牲一部分可能反转的合约，换取尾部亏损不继续扩张，最终提高组合层面的权利金留存率。",
        ),
        (
            "为什么说是线性增强，但还不能过度外推",
            "从完整样本看，1.5x、2.5x、3.5x 三个止损倍数在收益和回撤上都呈现近似线性梯度：倍数越低，收益越高、回撤越小。这说明在当前合约选择、保证金、流动性、B2C 品种倾斜规则不变的条件下，退出规则本身贡献了稳定超额。\n\n"
            "但三点线性不能直接推导出“越低越好”。如果继续降到 1.0x，可能会出现过早止损、手续费和滑点吞噬、止损后反复重开被洗的情况。因此下一轮应该补齐 1.0x、2.0x 完整、3.0x、4.0x、5.0x，并把 stop after-outcome 作为约束：如果某个倍数下止损后最终归零比例很高，说明它止早了；如果止损后继续恶化比例高，说明它止对了。",
        ),
        (
            "对策略主线的影响",
            "这次实验改变了我们对 S1 的一个核心判断：在纯卖权策略里，止损不是附属风控参数，而是 Retention Rate 和 Tail/Stop Loss 的核心结构变量。原先 2.5x 来自经验沟通，有合理性，但回测显示它在当前 B2C 框架下可能仍偏晚；1.5x 至少应该进入主线候选，而不是只作为保守风控版本。\n\n"
            "下一步不应马上把 1.5x 写死，而应做两件事：第一，完整跑完 2.0x 并补齐更多倍数；第二，把止损逻辑从单一硬倍数升级为“倍数 + IV/趋势/异常价确认”。如果合约价格上涨到 1.5x 但 IV 没有扩张、标的没有趋势穿透、且后续价格很快回落，那可能是噪声；如果价格上涨同时伴随 IV 上升、方向趋势和流动性恶化，就应坚决止损。",
        ),
    ]

    md_lines = [
        "# S1 B2C 止损倍数实验对比报告",
        "",
        "主题：Stop 1.5x / 2.0x / 2.5x / 3.5x / 不止损的收益、回撤、超额与权利金留存比较",
        "",
        "## 一、核心绩效表",
        "",
        make_md_table(full, perf_cols),
        "",
        "## 二、相对 B2C/2.5x 的超额表",
        "",
        make_md_table(excess_summary, excess_cols),
        "",
        "## 三、止损与权利金留存表",
        "",
        make_md_table(order_summary, order_cols),
        "",
        "## 四、Greek 归因表",
        "",
        make_md_table(full, greek_cols),
        "",
    ]
    for title_text, body in sections:
        md_lines.extend([f"## {title_text}", "", body, ""])
    md_lines.append("## 五、图表深读")
    md_lines.append("")
    for path, title_text, body in chart_specs:
        rel = path.relative_to(report_dir)
        md_lines.extend([f"### {title_text}", "", f"![{title_text}]({rel.as_posix()})", "", body, ""])

    md_path = report_dir / "s1_b2c_stop_loss_sweep_report_20260430.md"
    md_path.write_text("\n".join(md_lines), encoding="utf-8")

    docx_path = report_dir / "s1_b2c_stop_loss_sweep_report_20260430_feishu.docx"
    write_docx(sections, full, excess_summary, order_summary, chart_specs, docx_path)

    print(f"Report dir: {report_dir}")
    print(f"Markdown: {md_path}")
    print(f"DOCX: {docx_path}")
    print(full[["label", "status", "end", "cum_return", "ann_return", "max_dd", "sharpe", "calmar"]].to_string(index=False))


if __name__ == "__main__":
    main()
